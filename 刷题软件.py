import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import json, random, os, sys, hashlib, re, time
import urllib.request, urllib.error
import socket
try:
    import docx
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False
try:
    import openpyxl
    XLSX_AVAILABLE = True
except:
    XLSX_AVAILABLE = False
# ============ 多题库分离存储 ============
DATA_DIR = "D:\刷题软件\刷题数据"
USERS_FILE = os.path.join(DATA_DIR, "users.json")
ACCOUNTS_FILE = os.path.join(DATA_DIR, "accounts.json")
def ensure_data_dir():
    if not os.path.exists(DATA_DIR):
        os.makedirs(DATA_DIR)
# ============ 账号管理 ============
def load_accounts():
    ensure_data_dir()
    try:
        with open(ACCOUNTS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {}
def save_accounts(accounts):
    ensure_data_dir()
    with open(ACCOUNTS_FILE, "w", encoding="utf-8") as f:
        json.dump(accounts, f, ensure_ascii=False, indent=2)
def hash_password(pwd):
    return hashlib.sha256(pwd.encode()).hexdigest()
def is_valid_phone(phone):
    return re.match(r'^1[3-9]\d{9}$', phone) is not None
def is_valid_password(pwd):
    return len(pwd) >= 6

# ============ 网络请求 ============
SERVER_URL = ""
API_TIMEOUT = 5

def server_request(method, path, data=None):
    if not SERVER_URL:
        return None
    url = f"{SERVER_URL}/api{path}"
    try:
        body = json.dumps(data).encode() if data else None
        req = urllib.request.Request(url, data=body, method=method)
        req.add_header("Content-Type", "application/json")
        with urllib.request.urlopen(req, timeout=API_TIMEOUT) as resp:
            return json.loads(resp.read().decode())
    except urllib.error.URLError as e:
        return {"success": False, "msg": f"连接失败: {str(e.reason)}"}
    except socket.timeout:
        return {"success": False, "msg": "连接超时，请检查服务器是否开启"}
    except Exception as e:
        return {"success": False, "msg": f"请求失败: {str(e)}"}
def load_users():
    ensure_data_dir()
    try:
        with open(USERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except:
        return {"wrongNotes": [], "settings": {"shuffle": False, "showExp": True, "autoNext": False}}
def save_users(data):
    ensure_data_dir()
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
def load_bank(name):
    bank_file = os.path.join(DATA_DIR, f"{name}.json")
    try:
        with open(bank_file, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                # 清理非字典元素
                return [q for q in data if isinstance(q, dict)]
            return []
    except:
        return []
def save_bank(name, questions):
    ensure_data_dir()
    with open(os.path.join(DATA_DIR, f"{name}.json"), "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)
def list_banks():
    ensure_data_dir()
    files = os.listdir(DATA_DIR)
    banks = []
    for fname in files:
        if fname.endswith(".json") and fname not in ["users.json", "accounts.json", "remember.json"]:
            name = fname[:-5]
            qs = load_bank(name)
            # 清理混入的字符串元素
            qs = [q for q in qs if isinstance(q, dict)]
            total = len(qs)
            done = len([q for q in qs if q.get('done')])
            correct_cnt = len([q for q in qs if q.get('correct')])
            # 题型统计
            types = {}
            for q in qs:
                t = q.get('type', 'single')
                types[t] = types.get(t, 0) + 1
            # 正确率
            accuracy = round(correct_cnt / done * 100) if done > 0 else 0
            banks.append({"name": name, "total": total, "done": done,
                          "correct": correct_cnt, "acc": accuracy,
                          "types": types})
    return banks
def delete_bank(name):
    path = os.path.join(DATA_DIR, f"{name}.json")
    if os.path.exists(path):
        os.remove(path)
# ============ 登录界面 ============
class LoginWindow:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("刷题软件 - 登录")
        self.root.geometry("400x560")
        self.root.configure(bg="#1a1a2e")
        self.root.resizable(False, False)
        self.accounts = load_accounts()
        self.app = None
        self.remember_var = tk.BooleanVar(value=False)
        self.server_connected = False
        # 加载记住的账号和服务器
        remember_file = os.path.join(DATA_DIR, "remember.json")
        self.saved_phone = ""
        self.saved_pwd = ""
        self.saved_server = ""
        self.saved_role = "teacher"
        if os.path.exists(remember_file):
            try:
                with open(remember_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    self.saved_phone = data.get("phone", "")
                    self.saved_pwd = data.get("password", "")
                    self.saved_server = data.get("server", "")
                    self.saved_role = data.get("role", "teacher")
                    if self.saved_phone and self.saved_pwd:
                        self.remember_var.set(True)
            except:
                pass
        self.create_ui()
    def create_ui(self):
        # 标题
        tk.Label(self.root, text="🥟", font=("Microsoft YaHei", 36), bg="#1a1a2e", fg="#eee").pack(pady=(24, 0))
        tk.Label(self.root, text="刷题软件", font=("Microsoft YaHei", 16, "bold"), bg="#1a1a2e", fg="#e94560").pack(pady=(2, 6))
        # 表单容器
        self.login_frame = tk.Frame(self.root, bg="#1a1a2e")
        self.login_frame.pack(padx=40, fill="x")
        # 服务器地址
        tk.Label(self.login_frame, text="服务器地址", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        sv_frame = tk.Frame(self.login_frame, bg="#1a1a2e")
        sv_frame.pack(fill="x", pady=(0, 5))
        self.server_entry = tk.Entry(sv_frame, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                                    insertbackground="#eee", bd=0, relief="flat", highlightthickness=1,
                                    highlightbackground="#0f3460", highlightcolor="#e94560")
        self.server_entry.pack(side="left", fill="x", expand=True, ipady=5)
        self.server_entry.insert(0, self.saved_server)
        self.server_btn = tk.Button(sv_frame, text="连接", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#4ecca3",
                                   bd=0, padx=8, cursor="hand2", relief="flat", command=self.connect_server)
        self.server_btn.pack(side="left", padx=(4, 0))
        self.server_status = tk.Label(self.login_frame, text="", font=("Microsoft YaHei", 8), bg="#1a1a2e", fg="#888")
        self.server_status.pack(anchor="w", pady=(0, 4))
        # 角色选择
        role_frame = tk.Frame(self.login_frame, bg="#1a1a2e")
        role_frame.pack(fill="x", pady=(0, 5))
        self.role_var = tk.StringVar(value=self.saved_role if hasattr(self,'saved_role') else "teacher")
        tk.Radiobutton(role_frame, text="👨‍🏫 老师", variable=self.role_var, value="teacher",
                      font=("Microsoft YaHei", 10), bg="#1a1a2e", fg="#eee",
                      selectcolor="#1a1a2e", activebackground="#1a1a2e").pack(side="left", padx=6)
        tk.Radiobutton(role_frame, text="🎓 学员", variable=self.role_var, value="student",
                      font=("Microsoft YaHei", 10), bg="#1a1a2e", fg="#eee",
                      selectcolor="#1a1a2e", activebackground="#1a1a2e").pack(side="left", padx=6)
        # 手机号
        tk.Label(self.login_frame, text="手机号", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        self.phone_entry = tk.Entry(self.login_frame, font=("Microsoft YaHei", 13), bg="#16213e", fg="#eee",
                                   insertbackground="#eee", bd=0, relief="flat", highlightthickness=1,
                                   highlightbackground="#0f3460", highlightcolor="#e94560")
        self.phone_entry.pack(fill="x", ipady=6, pady=(0, 5))
        # 密码
        tk.Label(self.login_frame, text="密码", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        self.pwd_entry = tk.Entry(self.login_frame, font=("Microsoft YaHei", 13), bg="#16213e", fg="#eee",
                                 insertbackground="#eee", bd=0, relief="flat", highlightthickness=1,
                                 highlightbackground="#0f3460", highlightcolor="#e94560", show="*")
        self.pwd_entry.pack(fill="x", ipady=6, pady=(0, 5))
        self.pwd_entry.bind("<Return>", lambda e: self.do_login())
        # 提示
        self.msg_label = tk.Label(self.login_frame, text="", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#e94560")
        self.msg_label.pack(pady=(4, 2))
        # 记住我
        rem_frame = tk.Frame(self.login_frame, bg="#1a1a2e")
        rem_frame.pack(fill="x", pady=(2, 4))
        tk.Checkbutton(rem_frame, text="记住账号密码", variable=self.remember_var,
                      font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#aaa",
                      selectcolor="#1a1a2e", activebackground="#1a1a2e",
                      activeforeground="#aaa").pack(side="left")
        # 登录按钮
        self.login_btn = tk.Button(self.login_frame, text=" 登 录 ", font=("Microsoft YaHei", 12, "bold"),
                                  bg="#e94560", fg="#fff", bd=0, padx=16, pady=6,
                                  cursor="hand2", relief="flat", command=self.do_login)
        self.login_btn.pack(fill="x", ipady=3)
        # 注册
        tk.Label(self.login_frame, text="还没有账号？", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#666").pack(pady=(10, 0))
        tk.Button(self.login_frame, text="注 册", font=("Microsoft YaHei", 10), bg="#0f3460", fg="#4ecca3",
                 bd=0, padx=12, pady=3, cursor="hand2", relief="flat",
                 command=self.show_register).pack(pady=(3, 0))
        # 底部
        tk.Label(self.root, text="登录即表示同意使用条款", font=("Microsoft YaHei", 8), bg="#1a1a2e", fg="#555").pack(pady=(20, 10))
        # 填充已记住的账号密码
        if self.saved_phone:
            self.phone_entry.insert(0, self.saved_phone)
        if self.saved_pwd and hasattr(self, 'saved_role'):
            self.pwd_entry.insert(0, self.saved_pwd)
        # 自动连接服务器
        if self.saved_server:
            self.root.after(100, self.connect_server)
        self.root.mainloop()
    def connect_server(self):
        global SERVER_URL
        addr = self.server_entry.get().strip()
        if not addr:
            self.server_status.config(text="请输入服务器地址", fg="#e94560")
            return
        # 格式处理：补全 http://
        if not addr.startswith("http://") and not addr.startswith("https://"):
            addr = f"http://{addr}"
        SERVER_URL = addr.rstrip("/")
        self.server_status.config(text="正在连接...", fg="#f9c74f")
        self.root.update()
        result = server_request("GET", "/health")
        if result and result.get("status") == "ok":
            self.server_connected = True
            self.server_status.config(text="✅ 已连接", fg="#4ecca3")
            self.server_btn.config(text="已连接", state="disabled", bg="#4ecca3", fg="#1a1a2e")
        else:
            msg = result.get("msg", "无法连接") if result else "地址格式有误"
            self.server_status.config(text=f"❌ {msg}", fg="#e94560")
    def show_register(self):
        # 切换成注册模式 - 在同一窗口里切换 frame
        self.login_frame.pack_forget()
        self.show_register_form()
    def show_register_form(self):
        self.reg_frame = tk.Frame(self.root, bg="#1a1a2e")
        self.reg_frame.pack(padx=40, fill="x")
        tk.Label(self.reg_frame, text="注册新账号", font=("Microsoft YaHei", 15, "bold"), bg="#1a1a2e", fg="#e94560").pack(pady=(4, 6))
        tk.Label(self.reg_frame, text="手机号", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        self.reg_phone = tk.Entry(self.reg_frame, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                                   insertbackground="#eee", bd=0, highlightthickness=1,
                                   highlightbackground="#0f3460", highlightcolor="#e94560")
        self.reg_phone.pack(fill="x", ipady=5, pady=(0, 4))
        tk.Label(self.reg_frame, text="设置密码（至少6位）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        self.reg_pwd = tk.Entry(self.reg_frame, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                                 insertbackground="#eee", bd=0, highlightthickness=1,
                                 highlightbackground="#0f3460", highlightcolor="#e94560", show="*")
        self.reg_pwd.pack(fill="x", ipady=5, pady=(0, 4))
        tk.Label(self.reg_frame, text="确认密码", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(anchor="w")
        self.reg_pwd2 = tk.Entry(self.reg_frame, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                                  insertbackground="#eee", bd=0, highlightthickness=1,
                                  highlightbackground="#0f3460", highlightcolor="#e94560", show="*")
        self.reg_pwd2.pack(fill="x", ipady=5, pady=(0, 4))
        self.reg_msg = tk.Label(self.reg_frame, text="", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#e94560")
        self.reg_msg.pack(pady=(4, 2))
        tk.Button(self.reg_frame, text=" 注 册 ", font=("Microsoft YaHei", 12, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=16, pady=5,
                 cursor="hand2", relief="flat", command=self.do_register).pack(fill="x", ipady=3)
        tk.Button(self.reg_frame, text="← 返回登录", font=("Microsoft YaHei", 9), bg="#16213e", fg="#888",
                 bd=0, padx=10, pady=3, cursor="hand2", relief="flat",
                 command=self.show_login_form).pack(pady=(6, 0))
    def show_login_form(self):
        self.reg_frame.destroy()
        self.login_frame.pack(padx=40, fill="x")
    def do_register(self):
        phone = self.reg_phone.get().strip()
        pwd = self.reg_pwd.get()
        pwd2 = self.reg_pwd2.get()
        role = self.role_var.get()
        if not is_valid_phone(phone):
            self.reg_msg.config(text="请输入正确的11位手机号")
            return
        if not is_valid_password(pwd):
            self.reg_msg.config(text="密码至少6位")
            return
        if pwd != pwd2:
            self.reg_msg.config(text="两次密码不一致")
            return
        # 如果已连接服务器，走服务端注册
        if self.server_connected:
            result = server_request("POST", "/register", {"phone": phone, "password": pwd, "role": role})
            if result and result.get("success"):
                messagebox.showinfo("注册成功", "账号注册成功，请登录！")
                self.show_login_form()
            else:
                self.reg_msg.config(text=result.get("msg", "注册失败") if result else "连接失败")
            return
        # 本地注册
        accounts = load_accounts()
        if phone in accounts:
            self.reg_msg.config(text="该手机号已注册，请直接登录")
            return
        accounts[phone] = {"password": hash_password(pwd), "created_at": time.time(), "role": role}
        save_accounts(accounts)
        messagebox.showinfo("注册成功", "账号注册成功，请登录！")
        self.show_login_form()
    def do_login(self):
        phone = self.phone_entry.get().strip()
        pwd = self.pwd_entry.get()
        role = self.role_var.get()
        if not phone or not pwd:
            self.msg_label.config(text="请输入手机号和密码")
            return
        # 如果已连接服务器，走服务端登录
        if self.server_connected:
            result = server_request("POST", "/login", {"phone": phone, "password": pwd})
            if result and result.get("success"):
                server_role = result.get("role", "teacher")
                if server_role != role:
                    self.msg_label.config(text=f"该账号不是{role}角色")
                    return
            else:
                self.msg_label.config(text=result.get("msg", "登录失败") if result else "连接失败")
                return
        else:
            # 本地登录
            accounts = load_accounts()
            if phone not in accounts:
                self.msg_label.config(text="账号不存在，请先注册")
                return
            if accounts[phone]["password"] != hash_password(pwd):
                self.msg_label.config(text="密码错误")
                return
        # 记住账号密码
        remember_file = os.path.join(DATA_DIR, "remember.json")
        if self.remember_var.get():
            server_url = self.server_entry.get().strip()
            with open(remember_file, "w", encoding="utf-8") as f:
                json.dump({"phone": phone, "password": pwd, "server": server_url, "role": role}, f)
        else:
            try:
                if os.path.exists(remember_file):
                    os.remove(remember_file)
            except:
                pass
        self.root.destroy()
        self.app = QuizApp(tk.Tk(), phone, role)
        self.app.root.mainloop()
# ============ 主应用 ============
class QuizApp:
    def __init__(self, root, phone=None, role="teacher"):
        self.root = root
        self.root.title("刷题软件 v1.0")
        self.root.geometry("900x640")
        self.root.minsize(800, 580)
        self.root.configure(bg="#1a1a2e")
        self.phone = phone
        self.role = role
        self.users = load_users()
        self.wrong_notes = self.users.get("wrongNotes", [])
        self.settings = self.users.get("settings", {"shuffle": False, "showExp": True, "autoNext": False})
        self.current_bank = None
        self.current_type_filter = None
        self.practice_queue = []
        self.practice_idx = 0
        self.selected_opt = None
        self.nav_buttons = {}
        # 班级相关
        self.my_class_code = None
        self.my_class_name = ""
        self.class_banks = []
        self.setup_style()
        self.create_ui()
        self.show_home()
    def setup_style(self):
        s = ttk.Style()
        s.theme_use("clam")
        s.configure("TFrame", background="#1a1a2e")
        s.configure("Primary.TButton", background="#e94560", foreground="#fff", font=("Microsoft YaHei", 10, "bold"), padding=(10, 6))
        s.map("Primary.TButton", background=[("active", "#c73e54")])
        s.configure("Secondary.TButton", background="#0f3460", foreground="#fff", font=("Microsoft YaHei", 9), padding=(8, 5))
        s.map("Secondary.TButton", background=[("active", "#1a4a8a")])
        s.configure("Success.TButton", background="#4ecca3", foreground="#1a1a2e", font=("Microsoft YaHei", 9, "bold"), padding=(8, 5))
        s.map("Success.TButton", background=[("active", "#3db88a")])
    def create_ui(self):
        top = tk.Frame(self.root, bg="#16213e", height=52)
        top.pack(fill="x")
        top.pack_propagate(False)
        tk.Label(top, text="🥟 刷题软件", font=("Microsoft YaHei", 16, "bold"), bg="#16213e", fg="#e94560").pack(side="left", padx=20)
        self.stats_frame = tk.Frame(top, bg="#16213e")
        self.stats_frame.pack(side="right", padx=20)
        # 用户信息
        self.user_frame = tk.Frame(top, bg="#16213e")
        self.user_frame.pack(side="right", padx=(0, 4))
        if self.phone:
            tk.Label(self.user_frame, text=f"👤 {self.phone[:3]}****{self.phone[-4:]}",
                    font=("Microsoft YaHei", 9), bg="#16213e", fg="#4ecca3").pack(side="left", padx=2)
            tk.Button(self.user_frame, text="退出", font=("Microsoft YaHei", 9), bg="#16213e", fg="#e94560",
                     bd=0, padx=6, cursor="hand2", relief="flat",
                     command=self.logout).pack(side="left", padx=2)
        main = tk.Frame(self.root, bg="#1a1a2e")
        main.pack(fill="both", expand=True)
        sidebar = tk.Frame(main, bg="#16213e", width=180)
        sidebar.pack(side="left", fill="y")
        sidebar.pack_propagate(False)
        nav_items = [("🏠 首页", "home"), ("✍️ 题型练习", "type_practice"), ("📝 错题本", "wrong"),
                     ("📊 统计", "stats"), ("⚙️ 设置", "settings")]
        if self.role == "teacher":
            nav_items.insert(2, ("🏫 班级", "class"))
            nav_items.extend([("📚 题库管理", "banks"), ("📥 导入题库", "import")])
        for text, page in nav_items:
            btn = tk.Button(sidebar, text=text, font=("Microsoft YaHei", 10),
                          bg="#16213e", fg="#ccc", bd=0, padx=16, pady=10, anchor="w", cursor="hand2", relief="flat",
                          activebackground="#0f3460", activeforeground="#fff",
                          command=lambda p=page: self.navigate(p))
            btn.pack(fill="x")
            self.nav_buttons[page] = btn
        self.content_frame = tk.Frame(main, bg="#1a1a2e")
        self.content_frame.pack(side="right", fill="both", expand=True, padx=20, pady=16)
        self.update_top_stats()
    def update_top_stats(self):
        for w in self.stats_frame.winfo_children():
            w.destroy()
        banks = list_banks()
        total = sum(b['total'] for b in banks)
        done = sum(b['done'] for b in banks)
        correct = sum(b['correct'] for b in banks)
        acc = round(correct / done * 100) if done > 0 else 0
        for lbl, val in [("题库", len(banks)), ("总题数", total), ("正确率", f"{acc}%")]:
            tk.Label(self.stats_frame, text=f"{lbl} {val}", font=("Microsoft YaHei", 9), bg="#16213e", fg="#aaa").pack(side="left", padx=8)
    def navigate(self, page):
        self.current_page = page
        for btn in self.nav_buttons.values():
            btn.configure(bg="#16213e", fg="#ccc")
        if page in self.nav_buttons:
            self.nav_buttons[page].configure(bg="#0f3460", fg="#fff")
        for w in self.content_frame.winfo_children():
            w.destroy()
        if page == "home": self.show_home()
        elif page == "type_practice": self.show_type_practice()
        elif page == "wrong": self.show_wrong()
        elif page == "class": self.show_class()
        elif page == "banks": self.show_banks()
        elif page == "import": self.show_import()
        elif page == "stats": self.show_stats()
        elif page == "settings": self.show_settings()
    def clear_content(self):
        for w in self.content_frame.winfo_children():
            w.destroy()
    def card(self, parent=None):
        parent = parent or self.content_frame
        f = tk.Frame(parent, bg="#16213e")
        f.pack(fill="x", pady=6)
        return f
    def logout(self):
        if self.role == "student":
            self._save_student_progress()
        if messagebox.askyesno("退出登录", "确定要退出登录吗？"):
            self.root.destroy()
            LoginWindow()
    def get_all_questions(self):
        if self.current_bank:
            if str(self.current_bank).startswith("class_"):
                cache = getattr(self, 'class_bank_cache', {})
                return cache.get('questions', [])
            return load_bank(self.current_bank)
        all_q = []
        for b in list_banks():
            all_q.extend(load_bank(b['name']))
        return all_q
    # ===== 首页 =====
    def show_home(self):
        self.clear_content()
        tk.Label(self.content_frame, text=f"{'👨‍🏫 老师 - ' if self.role=='teacher' else '🎓 学员 - '}欢迎回来 🥟",
                font=("Microsoft YaHei", 22, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 16))
        if self.role == "student":
            # 学员首页：显示班级题库
            self._show_student_home()
        else:
            # 老师首页：显示本地题库
            self._show_teacher_home()
    def _show_teacher_home(self):
        banks = list_banks()
        total = sum(b['total'] for b in banks)
        done = sum(b['done'] for b in banks)
        correct = sum(b['correct'] for b in banks)
        acc = round(correct / done * 100) if done > 0 else 0
        grid = tk.Frame(self.content_frame, bg="#1a1a2e")
        grid.pack(fill="x", pady=(0, 14))
        for val, lbl, color in [(len(banks), "题库数量", "#e94560"), (total, "总题数", "#e94560"), (total - done, "未做", "#4ecca3"), (f"{acc}%", "正确率", "#f9c74f")]:
            c = tk.Frame(grid, bg="#16213e")
            c.pack(side="left", padx=5, fill="both", expand=True)
            tk.Label(c, text=str(val), font=("Microsoft YaHei", 24, "bold"), bg="#16213e", fg=color).pack(pady=(14, 4))
            tk.Label(c, text=lbl, font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(pady=(0, 12))
        type_grid = tk.Frame(self.content_frame, bg="#1a1a2e")
        type_grid.pack(fill="x", pady=(0, 14))
        for name, key, emoji, color in [("单选题", "single", "◻", "#4a90d9"), ("多选题", "multiple", "☑", "#e6a23c"), ("判断题", "judge", "✔", "#67c23a")]:
            qs = self.get_all_questions()
            cnt = len([q for q in qs if q.get('type') == key])
            if cnt == 0:
                continue
            tk.Button(type_grid, text=f"{emoji} {name}\n{cnt}题", font=("Microsoft YaHei", 11), bg=color, fg="#fff", bd=0, padx=14, pady=8, cursor="hand2", relief="flat",
                     command=lambda t=key: (setattr(self, 'current_bank', None), setattr(self, 'current_type_filter', t), self.navigate("type_practice"))).pack(side="left", padx=4)
        tk.Label(self.content_frame, text="📚 题库列表", font=("Microsoft YaHei", 14, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(6, 8))
        if not banks:
            tk.Label(self.content_frame, text="暂无题库，请导入题库后开始练习！", font=("Microsoft YaHei", 12), bg="#1a1a2e", fg="#888").pack(pady=40)
            tk.Label(self.content_frame, text="支持：Excel(.xlsx) / Word(.docx) / 文本(.txt)", font=("Microsoft YaHei", 10), bg="#1a1a2e", fg="#666").pack()
            return
        for b in banks:
            bc = tk.Frame(self.content_frame, bg="#16213e")
            bc.pack(fill="x", pady=4)
            color = b['types'].get('judge') and "#67c23a" or b['types'].get('multiple') and "#e6a23c" or "#4a90d9"
            tk.Frame(bc, bg=color, width=4).pack(side="left", fill="y")
            inner = tk.Frame(bc, bg="#16213e")
            inner.pack(side="left", fill="both", expand=True, padx=12, pady=8)
            names = {'single': '单选', 'multiple': '多选', 'judge': '判断', 'fill': '填空', 'essay': '问答'}
            tk.Label(inner, text=f"📚 {b['name']}", font=("Microsoft YaHei", 12, "bold"), bg="#16213e", fg="#eee").pack(anchor="w")
            tk.Label(inner, text=" | ".join([f"{names.get(t,t)} {c}题" for t, c in b['types'].items()]), font=("Microsoft YaHei", 9), bg="#16213e", fg="#888").pack(anchor="w", pady=(2, 6))
            prog_bg = tk.Frame(inner, bg="#0f3460", height=6)
            prog_bg.pack(fill="x", pady=(0, 6))
            prog_bg.pack_propagate(False)
            if b['done'] > 0:
                tk.Frame(prog_bg, bg="#e94560", height=6).place(x=0, y=0, relwidth=b['done'] / b['total'], relheight=1)
            btn_row = tk.Frame(inner, bg="#16213e")
            btn_row.pack(anchor="e")
            for lbl, cmd in [("练习", lambda n=b['name']: (setattr(self, 'current_bank', n), setattr(self, 'current_type_filter', None), self.navigate("type_practice"))),
                             ("统计", lambda n=b['name']: (setattr(self, 'current_bank', n), self.navigate("stats")))]:
                tk.Button(btn_row, text=f"✍️ {lbl}" if lbl == "练习" else f"📊 {lbl}", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#aaa", bd=0, padx=8, pady=3, cursor="hand2", relief="flat", command=cmd).pack(side="left", padx=2)
            tk.Button(btn_row, text="🗑️", font=("Microsoft YaHei", 9), bg="#16213e", fg="#e94560", bd=0, padx=8, cursor="hand2", relief="flat",
                     command=lambda n=b['name']: (delete_bank(n), self.show_home(), self.update_top_stats())).pack(side="left", padx=2)
        imp = self.card()
        tk.Label(imp, text="➕ 导入新题库", font=("Microsoft YaHei", 11, "bold"), bg="#16213e", fg="#4ecca3").pack(anchor="w", padx=14, pady=10)
        tk.Button(imp, text="  📂 从文档导入（Excel/Word/TXT）", font=("Microsoft YaHei", 10), bg="#16213e", fg="#aaa", bd=0, padx=14, pady=8, cursor="hand2", relief="flat",
                 command=lambda: self.navigate("import")).pack(anchor="w")
    def _show_student_home(self):
        """学员首页：仅显示已购买的班级公开题库"""
        if not SERVER_URL:
            tk.Label(self.content_frame, text="未连接到服务器，请重新登录并连接服务器",
                    font=("Microsoft YaHei", 13), bg="#1a1a2e", fg="#e94560").pack(pady=40)
            return
        result = server_request("POST", "/get_student_classes", {"student": self.phone})
        if not result or not result.get("success"):
            tk.Label(self.content_frame, text="获取班级信息失败", font=("Microsoft YaHei", 13),
                    bg="#1a1a2e", fg="#e94560").pack(pady=40)
            return
        classes = result.get("classes", [])
        if not classes:
            tk.Label(self.content_frame, text="你还没有加入任何班级", font=("Microsoft YaHei", 14),
                    bg="#1a1a2e", fg="#888").pack(pady=30)
            tk.Button(self.content_frame, text="🏫 加入班级", font=("Microsoft YaHei", 12, "bold"),
                     bg="#e94560", fg="#fff", bd=0, padx=20, pady=8,
                     cursor="hand2", relief="flat",
                     command=lambda: self.navigate("class")).pack(pady=8)
            return
        # 显示所有班级题库
        all_banks = []
        for cls_info in classes:
            for b in cls_info.get("banks", []):
                b["class"] = cls_info["name"]
                all_banks.append(b)
        total = sum(b.get('count', 0) for b in all_banks)
        grid = tk.Frame(self.content_frame, bg="#1a1a2e")
        grid.pack(fill="x", pady=(0, 10))
        for val, lbl, color in [(len(classes), "已加入班级", "#4ecca3"), (len(all_banks), "可练题库", "#4a90d9"), (total, "总题数", "#e94560")]:
            c = tk.Frame(grid, bg="#16213e")
            c.pack(side="left", padx=5, fill="both", expand=True)
            tk.Label(c, text=str(val), font=("Microsoft YaHei", 24, "bold"), bg="#16213e", fg=color).pack(pady=(14, 4))
            tk.Label(c, text=lbl, font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(pady=(0, 12))
        tk.Label(self.content_frame, text="📚 班级题库", font=("Microsoft YaHei", 14, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(6, 8))
        for cls_info in classes:
            card = self.card()
            tk.Label(card, text=f"🏫 {cls_info['name']}", font=("Microsoft YaHei", 12, "bold"),
                    bg="#16213e", fg="#eee").pack(anchor="w", padx=14, pady=(8, 4))
            if not cls_info.get("banks"):
                tk.Label(card, text="  暂无题库", font=("Microsoft YaHei", 10), bg="#16213e", fg="#666").pack(anchor="w", padx=14, pady=(0, 8))
            for b in cls_info.get("banks", []):
                bf = tk.Frame(card, bg="#0f3460")
                bf.pack(fill="x", padx=14, pady=2)
                tk.Label(bf, text=f"📖 {b['name']}  ({b['count']}题)", font=("Microsoft YaHei", 10),
                        bg="#0f3460", fg="#ccc").pack(side="left", padx=10, pady=6)
                tk.Button(bf, text="练习", font=("Microsoft YaHei", 9), bg="#4ecca3", fg="#1a1a2e",
                         bd=0, padx=10, pady=2, cursor="hand2", relief="flat",
                         command=lambda n=b['name']: (setattr(self, 'current_bank', f"class_{n}"), self.load_class_bank(n), self.navigate("type_practice"))
                         ).pack(side="right", padx=6, pady=4)
    # ===== 班级页面（老师管理班级，学员加入班级） =====
    def show_class(self):
        self.clear_content()
        if not SERVER_URL:
            tk.Label(self.content_frame, text="未连接到服务器", font=("Microsoft YaHei", 14),
                    bg="#1a1a2e", fg="#e94560").pack(pady=40)
            return
        if self.role == "teacher":
            self._show_teacher_class()
        else:
            self._show_student_class()
    def _show_teacher_class(self):
        tk.Label(self.content_frame, text="🏫 班级管理", font=("Microsoft YaHei", 18, "bold"),
                bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 12))
        c = self.card()
        tk.Label(c, text="创建新班级", font=("Microsoft YaHei", 12, "bold"),
                bg="#16213e", fg="#4ecca3").pack(anchor="w", padx=14, pady=(8, 2))
        name_row = tk.Frame(c, bg="#16213e")
        name_row.pack(fill="x", padx=14, pady=(0, 8))
        self.class_name_entry = tk.Entry(name_row, font=("Microsoft YaHei", 11), bg="#0f3460", fg="#eee",
                                        insertbackground="#eee", bd=0, highlightthickness=1,
                                        highlightbackground="#0f3460")
        self.class_name_entry.pack(side="left", fill="x", expand=True, ipady=4, padx=(0, 6))
        self.class_name_entry.insert(0, "新班级")
        tk.Button(name_row, text="创建", font=("Microsoft YaHei", 10, "bold"), bg="#e94560", fg="#fff",
                 bd=0, padx=14, pady=5, cursor="hand2", relief="flat",
                 command=self._do_create_class).pack(side="left")
        # 显示已有班级
        result = server_request("POST", "/get_teacher_classes", {"teacher": self.phone})
        if result and result.get("success"):
            for cls_info in result.get("classes", []):
                card = self.card()
                # 班级头部
                hdr = tk.Frame(card, bg="#16213e")
                hdr.pack(fill="x", padx=14, pady=(8, 2))
                tk.Label(hdr, text=f"🏫 {cls_info['name']}", font=("Microsoft YaHei", 13, "bold"),
                        bg="#16213e", fg="#eee").pack(side="left")
                tk.Label(hdr, text=f"  学员 {cls_info['student_count']}人", font=("Microsoft YaHei", 10),
                        bg="#16213e", fg="#aaa").pack(side="left", padx=4)
                tk.Label(card, text=f" 加入码: {cls_info['code']}", font=("Microsoft YaHei", 11, "bold"),
                        bg="#16213e", fg="#f9c74f").pack(anchor="w", padx=14, pady=(0, 4))
                # 已上传题库
                tk.Label(card, text="已上传题库:", font=("Microsoft YaHei", 9), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(2, 2))
                for b in cls_info.get("banks", []):
                    bf = tk.Frame(card, bg="#0f3460")
                    bf.pack(fill="x", padx=14, pady=2)
                    tk.Label(bf, text=f"📖 {b['name']}  ({b['count']}题)", font=("Microsoft YaHei", 10),
                            bg="#0f3460", fg="#ccc").pack(side="left", padx=10, pady=4)
                    tk.Button(bf, text="移除", font=("Microsoft YaHei", 8), bg="#e94560", fg="#fff",
                             bd=0, padx=8, pady=1, cursor="hand2", relief="flat",
                             command=lambda bn=b['name'], cc=cls_info['code']: (self._remove_class_bank(cc, bn), self.show_class())
                             ).pack(side="right", padx=6, pady=2)
                # 上传题库按钮
                btn_row = tk.Frame(card, bg="#16213e")
                btn_row.pack(padx=14, pady=(6, 8))
                tk.Button(btn_row, text="📤 上传本地题库到班级", font=("Microsoft YaHei", 9, "bold"),
                         bg="#4ecca3", fg="#1a1a2e", bd=0, padx=10, pady=4,
                         cursor="hand2", relief="flat",
                         command=lambda cc=cls_info['code']: self._upload_bank_to_class(cc)
                         ).pack(side="left", padx=2)
    def _show_student_class(self):
        tk.Label(self.content_frame, text="🏫 加入班级", font=("Microsoft YaHei", 18, "bold"),
                bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 12))
        # 加入班级
        c = self.card()
        tk.Label(c, text="输入班级加入码", font=("Microsoft YaHei", 12), bg="#16213e", fg="#4ecca3").pack(anchor="w", padx=14, pady=(8, 4))
        row = tk.Frame(c, bg="#16213e")
        row.pack(fill="x", padx=14, pady=(0, 10))
        self.join_code_entry = tk.Entry(row, font=("Microsoft YaHei", 16, "bold"), bg="#0f3460", fg="#eee",
                                       insertbackground="#eee", bd=0, highlightthickness=2,
                                       highlightbackground="#e94560", justify="center", width=8)
        self.join_code_entry.pack(side="left", ipady=6, padx=(0, 10))
        self.join_code_entry.bind("<Return>", lambda e: self._do_join_class())
        tk.Button(row, text="加入", font=("Microsoft YaHei", 12, "bold"), bg="#4ecca3", fg="#1a1a2e",
                 bd=0, padx=20, pady=6, cursor="hand2", relief="flat",
                 command=self._do_join_class).pack(side="left")
        self.join_msg = tk.Label(c, text="", font=("Microsoft YaHei", 9), bg="#16213e", fg="#e94560")
        self.join_msg.pack(anchor="w", padx=14)
        # 已加入的班级
        result = server_request("POST", "/get_student_classes", {"student": self.phone})
        if result and result.get("success"):
            for cls_info in result.get("classes", []):
                card = self.card()
                tk.Label(card, text=f"🏫 {cls_info['name']}", font=("Microsoft YaHei", 13, "bold"),
                        bg="#16213e", fg="#eee").pack(anchor="w", padx=14, pady=(8, 4))
                for b in cls_info.get("banks", []):
                    bf = tk.Frame(card, bg="#0f3460")
                    bf.pack(fill="x", padx=14, pady=2)
                    tk.Label(bf, text=f"📖 {b['name']}  ({b['count']}题)", font=("Microsoft YaHei", 10),
                            bg="#0f3460", fg="#ccc").pack(side="left", padx=10, pady=4)
                    tk.Button(bf, text="练习", font=("Microsoft YaHei", 9), bg="#4ecca3", fg="#1a1a2e",
                             bd=0, padx=10, pady=2, cursor="hand2", relief="flat",
                             command=lambda n=b['name']: (setattr(self, 'current_bank', f"class_{n}"), self.load_class_bank(n), self.navigate("type_practice"))
                             ).pack(side="right", padx=6, pady=4)
    def _do_create_class(self):
        name = self.class_name_entry.get().strip()
        if not name:
            name = "新班级"
        result = server_request("POST", "/create_class", {"teacher": self.phone, "name": name})
        if result and result.get("success"):
            messagebox.showinfo("成功", f"班级创建成功！\n加入码: {result['code']}\n\n把加入码发给学员即可加入")
            self.show_class()
        else:
            messagebox.showerror("失败", result.get("msg", "创建失败") if result else "连接失败")
    def _upload_bank_to_class(self, class_code):
        banks = list_banks()
        if not banks:
            messagebox.showwarning("提示", "本地没有题库可上传")
            return
        # 弹窗选择题库
        win = tk.Toplevel(self.root)
        win.title("选择题库上传")
        win.geometry("360x300")
        win.configure(bg="#1a1a2e")
        win.transient(self.root)
        win.grab_set()
        tk.Label(win, text="选择题库上传到班级", font=("Microsoft YaHei", 12, "bold"),
                bg="#1a1a2e", fg="#eee").pack(pady=12)
        listbox = tk.Listbox(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                            bd=0, highlightthickness=0, selectbackground="#e94560")
        listbox.pack(fill="both", expand=True, padx=20, pady=6)
        for b in banks:
            listbox.insert(tk.END, f"{b['name']}  ({b['total']}题)")
        def do_upload():
            sel = listbox.curselection()
            if not sel:
                return
            bank_name = banks[sel[0]]['name']
            result = server_request("POST", "/upload_bank", {"teacher": self.phone, "bank_name": bank_name, "class_code": class_code})
            if result and result.get("success"):
                messagebox.showinfo("成功", "题库已上传到班级！")
                win.destroy()
                self.show_class()
            else:
                messagebox.showerror("失败", result.get("msg", "上传失败") if result else "连接失败")
        tk.Button(win, text="上传", font=("Microsoft YaHei", 11, "bold"), bg="#e94560", fg="#fff",
                 bd=0, padx=16, pady=5, cursor="hand2", relief="flat",
                 command=do_upload).pack(pady=8)
    def _remove_class_bank(self, class_code, bank_name):
        result = server_request("POST", "/remove_class_bank", {"teacher": self.phone, "bank_name": bank_name, "class_code": class_code})
        if result and result.get("success"):
            self.show_class()
    def _do_join_class(self):
        code = self.join_code_entry.get().strip()
        if not code:
            self.join_msg.config(text="请输入加入码")
            return
        result = server_request("POST", "/join_class", {"student": self.phone, "code": code})
        if result and result.get("success"):
            self.join_msg.config(text=f"✅ 已加入{result.get('class_name', '班级')}！", fg="#4ecca3")
            self.root.after(1000, self.show_class)
        else:
            self.join_msg.config(text=result.get("msg", "加入失败") if result else "连接失败", fg="#e94560")
    def load_class_bank(self, bank_name):
        """从服务器加载班级题库到本地缓存"""
        result = server_request("POST", "/get_bank_questions", {"bank_name": bank_name})
        if result and result.get("success"):
            self.class_bank_cache = {"name": f"class_{bank_name}", "questions": result.get("questions", [])}
    # ===== 题型练习 =====
    def show_type_practice(self):
        self.clear_content()
        # 判断是否有班级题库
        if self.current_bank and str(self.current_bank).startswith("class_"):
            banks = [{"name": self.current_bank, "total": len(getattr(self, 'class_bank_cache', {}).get('questions', [])),
                      "types": {}, "done": 0, "correct": 0, "acc": 0}]
        else:
            banks = list_banks()
        if not banks:
            tk.Label(self.content_frame, text="暂无题库，请先导入题库！", font=("Microsoft YaHei", 13), bg="#1a1a2e", fg="#888").pack(pady=40)
            ttk.Button(self.content_frame, text="去导入题库", style="Primary.TButton", command=lambda: self.navigate("import")).pack()
            return
        tk.Label(self.content_frame, text="✍️ 题型练习", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 12))
        c = self.card()
        tk.Label(c, text="选择题库", font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(8, 4))
        bsel = tk.Frame(c, bg="#16213e")
        bsel.pack(fill="x", padx=12, pady=(0, 8))
        is_all = self.current_bank is None
        tk.Button(bsel, text="全部题库", font=("Microsoft YaHei", 10), bg="#0f3460" if is_all else "#16213e", fg="#fff" if is_all else "#888", bd=0, padx=12, pady=6, cursor="hand2", relief="flat",
                 command=lambda: (setattr(self, 'current_bank', None), self.show_type_practice())).pack(side="left", padx=2)
        for b in banks:
            is_sel = self.current_bank == b['name']
            tk.Button(bsel, text=f"{b['name']}({b['total']}题)", font=("Microsoft YaHei", 10), bg="#0f3460" if is_sel else "#16213e",
                     fg="#fff" if is_sel else "#888", bd=0, padx=10, pady=6, cursor="hand2", relief="flat",
                     command=lambda n=b['name']: (setattr(self, 'current_bank', n), self.show_type_practice())).pack(side="left", padx=2)
        c2 = self.card()
        tk.Label(c2, text="选择题型", font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(8, 4))
        tsel = tk.Frame(c2, bg="#16213e")
        tsel.pack(fill="x", padx=12, pady=(0, 8))
        qs = self.get_all_questions()
        type_info = [(None, "全部", "#e94560"), ("single", "单选题", "#4a90d9"), ("multiple", "多选题", "#e6a23c"),
                      ("judge", "判断题", "#67c23a"), ("fill", "填空题", "#9c27b0"), ("essay", "问答题", "#795548")]
        for key, label, color in type_info:
            cnt = len([q for q in qs if q.get('type') == key]) if key else len(qs)
            if key and cnt == 0:
                continue
            is_sel = self.current_type_filter == key
            tk.Button(tsel, text=f"{label}({cnt})", font=("Microsoft YaHei", 10),
                     bg=color if is_sel else "#0f3460", fg="#fff", bd=0, padx=12, pady=6, cursor="hand2", relief="flat",
                     command=lambda t=key: (setattr(self, 'current_type_filter', t), self.show_type_practice())).pack(side="left", padx=2)
        start = self.card()
        filtered = [q for q in qs if not self.current_type_filter or q.get('type') == self.current_type_filter]
        if not filtered:
            tk.Label(start, text="该条件下没有可练习的题目！", font=("Microsoft YaHei", 12), bg="#16213e", fg="#888").pack(padx=14, pady=20)
            return
        tk.Label(start, text=f"共 {len(filtered)} 道题目", font=("Microsoft YaHei", 11), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(6, 4))
        if self.role == "student" and self.current_bank and str(self.current_bank).startswith("class_"):
            real_name = self.current_bank.replace("class_", "", 1)
            progress_res = server_request("POST", "/load_bank_progress", {"phone": self.phone, "bank_name": real_name})
            if progress_res and progress_res.get("success") and progress_res.get("questions"):
                # 用服务端进度合并题库
                progress_qs = progress_res["questions"]
                progress_map = {q.get('id'): q for q in progress_qs if q.get('id')}
                for q in filtered:
                    qid = q.get('id')
                    if qid and qid in progress_map:
                        q['done'] = progress_map[qid].get('done', False)
                        q['correct'] = progress_map[qid].get('correct', False)
            # 加载错题本
            wn_res = server_request("POST", "/load_wrong_notes", {"phone": self.phone})
            if wn_res and wn_res.get("success"):
                self.wrong_notes = wn_res.get("wrong_notes", [])
                self.users['wrongNotes'] = self.wrong_notes
        tk.Button(start, text="▶️ 开始练习", font=("Microsoft YaHei", 13, "bold"), bg="#e94560", fg="#fff", bd=0, padx=24, pady=10, cursor="hand2", relief="flat",
                 command=lambda: (setattr(self, 'settings', load_users().get('settings', self.settings)), setattr(self, 'practice_queue', random.sample(filtered, len(filtered)) if self.settings.get('shuffle') else filtered[:]), setattr(self, 'practice_idx', 0), self.show_practice_question())).pack(pady=(6, 10))
    # ===== 答题界面 =====
    def show_practice_question(self):
        self.clear_content()
        if self.practice_idx >= len(self.practice_queue):
            self.show_practice_complete()
            return
        q = self.practice_queue[self.practice_idx]
        self.selected_opt = None
        total_q = len(self.practice_queue)
        
        # 整体两栏布局：左答题卡 + 右题目区
        main_row = tk.Frame(self.content_frame, bg="#1a1a2e")
        main_row.pack(fill="both", expand=True)
        
        # === 左栏：答题卡（固定窄宽度） ===
        left = tk.Frame(main_row, bg="#16213e", width=220)
        left.pack(side="left", fill="y")
        left.pack_propagate(False)
        
        tk.Label(left, text="答题卡", font=("Microsoft YaHei", 9, "bold"),
                bg="#16213e", fg="#888").pack(pady=(8, 4))
        
        # 统计
        done_count = sum(1 for qq in self.practice_queue if isinstance(qq, dict) and qq.get('done'))
        correct_count = sum(1 for qq in self.practice_queue if isinstance(qq, dict) and qq.get('correct'))
        tk.Label(left, text=f"已答 {done_count}/{total_q}", font=("Microsoft YaHei", 7),
                bg="#16213e", fg="#aaa").pack()
        tk.Label(left, text=f"正确 {correct_count}", font=("Microsoft YaHei", 7),
                bg="#16213e", fg="#4ecca3").pack(pady=(0, 6))
        
        # 答题卡片网格（可滚动）
        card_frame = tk.Frame(left, bg="#16213e")
        card_frame.pack(fill="both", expand=True, padx=4, pady=0)
        up_frame = tk.Frame(card_frame, bg="#16213e")
        up_frame.pack(fill="both", expand=True)
        card_canvas = tk.Canvas(up_frame, bg="#16213e", bd=0, highlightthickness=0, width=175)
        card_scroll = tk.Scrollbar(up_frame, orient="vertical", command=card_canvas.yview, width=14)
        card_canvas.configure(yscrollcommand=card_scroll.set)
        inner_frame = tk.Frame(card_canvas, bg="#16213e")
        inner_frame.bind("<Configure>", lambda e: card_canvas.configure(scrollregion=card_canvas.bbox("all")))
        card_canvas.create_window((0, 0), window=inner_frame, anchor="nw")
        def _on_card_wheel(event):
            card_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        card_canvas.bind("<Enter>", lambda e: card_canvas.bind_all("<MouseWheel>", _on_card_wheel))
        card_canvas.bind("<Leave>", lambda e: card_canvas.unbind_all("<MouseWheel>"))
        card_canvas.pack(side="left", fill="both", expand=True)
        card_scroll.pack(side="right", fill="y")
        
        per_col = 5
        for j in range(total_q):
            if j % per_col == 0:
                rowf = tk.Frame(inner_frame, bg="#16213e")
                rowf.pack(pady=1)
            is_current = j == self.practice_idx
            qq = self.practice_queue[j] if isinstance(self.practice_queue[j], dict) else {}
            if qq.get('done'):
                if qq.get('correct'):
                    bg_c = "#4ecca3"
                    fg_c = "#1a1a2e"
                else:
                    bg_c = "#e94560"
                    fg_c = "#fff"
            else:
                bg_c = "#0f3460" if not is_current else "#4a90d9"
                fg_c = "#aaa" if not is_current else "#fff"
            tk.Button(rowf, text=str(j+1), font=("Microsoft YaHei", 8),
                     bg=bg_c, fg=fg_c, bd=0, width=3, height=1,
                     cursor="hand2", relief="flat",
                     command=lambda idx=j: (setattr(self, 'practice_idx', idx), self.show_practice_question())
                     ).pack(side="left", padx=1, pady=0)
        
        # 左侧底部操作跳题
        tk.Label(left, text="跳至", font=("Microsoft YaHei", 8), bg="#16213e", fg="#666").pack(pady=(12, 2))
        sv = tk.StringVar(value=str(self.practice_idx + 1))
        tk.Spinbox(left, from_=1, to=total_q, textvariable=sv,
                 width=3, font=("Microsoft YaHei", 8),
                 bg="#0f3460", fg="#eee", bd=0, buttonbackground="#16213e",
                 readonlybackground="#0f3460", justify="center").pack(ipady=2)
        tk.Button(left, text="跳转", font=("Microsoft YaHei", 8),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=10, pady=1,
                 cursor="hand2", relief="flat",
                 command=lambda: (setattr(self, 'practice_idx', max(0, min(int(sv.get() or '1') - 1, total_q - 1))), self.show_practice_question())
                 ).pack(pady=(2, 8))
        
        tk.Button(left, text="返回", font=("Microsoft YaHei", 8),
                 bg="#0f3460", fg="#aaa", bd=0, padx=10, pady=1,
                 cursor="hand2", relief="flat",
                 command=lambda: (setattr(self, 'practice_idx', 0), self.show_type_practice())
                 ).pack(pady=(0, 4))
        
        # === 右栏：题目 + 选项 ===
        right = tk.Frame(main_row, bg="#1a1a2e")
        right.pack(side="right", fill="both", expand=True)
        
        # 进度条
        prog = tk.Frame(right, bg="#0f3460", height=6)
        prog.pack(fill="x")
        prog.pack_propagate(False)
        tk.Frame(prog, bg="#e94560").place(x=0, y=0, relwidth=(self.practice_idx+1) / total_q, relheight=1)
        
        # 题号 + 类型标签行
        info_row = tk.Frame(right, bg="#1a1a2e")
        info_row.pack(fill="x", pady=(8, 4), padx=14)
        tag_colors = {'single': '#4a90d9', 'multiple': '#e6a23c', 'judge': '#67c23a', 'fill': '#9c27b0', 'essay': '#795548'}
        tag_names = {'single': '单选题', 'multiple': '多选题', 'judge': '判断题', 'fill': '填空题', 'essay': '问答题'}
        tk.Label(info_row, text=f"第{self.practice_idx+1}题", font=("Microsoft YaHei", 11, "bold"),
                bg="#1a1a2e", fg="#eee").pack(side="left")
        # 编辑按钮
        tk.Button(info_row, text="✏️ 编辑", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#f9c74f",
                 bd=0, padx=8, pady=2, cursor="hand2", relief="flat",
                 command=lambda: self.edit_current_question()).pack(side="left", padx=(10, 0))
        tk.Label(info_row, text=f"  {tag_names.get(q.get('type'), '')}  ",
                font=("Microsoft YaHei", 8),
                bg=tag_colors.get(q.get('type'), '#888'), fg="#fff",
                padx=6, pady=1).pack(side="left", padx=(8, 0))
        
        # 题干 + 选项（固定区域，无需滚动）
        c = tk.Frame(right, bg="#16213e", bd=1, relief="flat")
        c.pack(fill="both", expand=True, padx=14, pady=(0, 6))
        
        tk.Label(c, text=q.get('text', ''), font=("Microsoft YaHei", 13),
                bg="#16213e", fg="#eee", wraplength=550, anchor="w", justify="left"
                ).pack(anchor="w", padx=16, pady=(16, 12))
        
        # 选项列表（考试宝风格：横条按钮）
        self.opt_buttons = []
        self.opt_vars = [tk.StringVar(value='') for _ in q.get('options', [])]
        for i, opt in enumerate(q.get('options', [])):
            letter = chr(65+i)
            opt_item = tk.Frame(c, bg="#1a1a2e", bd=1, relief="flat", highlightthickness=1, highlightbackground="#0f3460")
            opt_item.pack(fill="x", padx=16, pady=4)
            
            # 字母圆圈
            letter_lbl = tk.Label(opt_item, text=letter, font=("Microsoft YaHei", 12, "bold"),
                                 bg="#0f3460", fg="#eee", width=3, height=1, padx=4, pady=4)
            letter_lbl.pack(side="left", padx=(4, 8), pady=4)
            
            # 选项文本
            txt_lbl = tk.Label(opt_item, text=opt.get('text', ''), font=("Microsoft YaHei", 12),
                              bg="#1a1a2e", fg="#ddd", anchor="w", justify="left", wraplength=460)
            txt_lbl.pack(side="left", fill="x", expand=True, padx=(0, 8), pady=4)
            
            # 点击整个条目
            def make_click(idx):
                def click():
                    self.select_practice_opt(idx)
                return click
            opt_item.bind("<Button-1>", lambda e, idx=i: self.select_practice_opt(idx))
            letter_lbl.bind("<Button-1>", lambda e, idx=i: self.select_practice_opt(idx))
            txt_lbl.bind("<Button-1>", lambda e, idx=i: self.select_practice_opt(idx))
            
            opt_item.letter_lbl = letter_lbl
            opt_item.txt_lbl = txt_lbl
            self.opt_buttons.append(opt_item)
        
        # 解析
        self.exp_label = tk.Label(c, text="", font=("Microsoft YaHei", 10), bg="#16213e", fg="#f9c74f",
                                  wraplength=580, anchor="w", justify="left")
        self.exp_label.pack(fill="x", padx=16, pady=(8, 0))
        self.exp_label.pack_forget()
        
        # 底部操作按钮
        act = tk.Frame(right, bg="#1a1a2e")
        act.pack(fill="x", padx=14, pady=(4, 8))
        self.submit_btn = tk.Button(act, text="提交答案", font=("Microsoft YaHei", 11, "bold"),
                                   bg="#4ecca3", fg="#1a1a2e", bd=0, padx=20, pady=6,
                                   cursor="hand2", relief="flat",
                                   command=self.submit_practice_answer)
        self.submit_btn.pack(side="left")
        self.next_btn = tk.Button(act, text="下一题 →", font=("Microsoft YaHei", 11, "bold"),
                                 bg="#4a90d9", fg="#fff", bd=0, padx=20, pady=6,
                                 cursor="hand2", relief="flat",
                                 command=self.next_practice)
        self.next_btn.pack(side="left", padx=(8, 0))
        self.next_btn.pack_forget()

    def select_practice_opt(self, idx):
        q = self.practice_queue[self.practice_idx] if self.practice_idx < len(self.practice_queue) else None
        if not q:
            return
        opts = q.get('options', [])
        q_type = q.get('type', 'single')
        is_multi = q_type == 'multiple'
        
        was_selected = self.opt_vars[idx].get() == 'selected'
        
        def _set_opt(i, selected):
            item = self.opt_buttons[i]
            bg = "#0f3460"
            fg = "#eee"
            letter_bg = "#0f3460"
            if selected:
                bg = "#1a4a3e" if not is_multi else "#3a3a1a"
                fg = "#4ecca3" if not is_multi else "#e6a23c"
                letter_bg = "#4ecca3" if not is_multi else "#e6a23c"
            else:
                if q.get('done') and opts[i].get('correct'):
                    bg = "#1a4a3e"
                    fg = "#4ecca3"
                    letter_bg = "#4ecca3"
            item.config(bg=bg, highlightbackground=bg)
            if hasattr(item, 'letter_lbl'):
                item.letter_lbl.config(bg=letter_bg, fg="#1a1a2e" if selected else "#eee")
            if hasattr(item, 'txt_lbl'):
                item.txt_lbl.config(bg=bg, fg=fg)
        
        if not is_multi:
            for i in range(len(opts)):
                self.opt_vars[i].set('')
                _set_opt(i, False)
            if not was_selected:
                self.opt_vars[idx].set('selected')
                _set_opt(idx, True)
        else:
            _set_opt(idx, not was_selected)
            self.opt_vars[idx].set('' if was_selected else 'selected')
    def submit_practice_answer(self):
        q = self.practice_queue[self.practice_idx]
        opts = q.get('options', [])
        q_type = q.get('type', 'single')
        is_multi = q_type == 'multiple'
        selected_indices = [i for i in range(len(opts)) if self.opt_vars[i].get()]
        if not selected_indices:
            messagebox.showwarning("提示", "请至少选择一个选项！")
            return
        if not is_multi and len(selected_indices) > 1:
            messagebox.showwarning("提示", "单选题只能选择一个选项！")
            return
        if is_multi:
            correct_indices = [i for i, o in enumerate(opts) if o.get('correct')]
            is_correct = set(selected_indices) == set(correct_indices)
        else:
            is_correct = opts[selected_indices[0]].get('correct', False)
        q['done'] = True
        q['correct'] = is_correct
        # 显示正确答案（绿色）和错误选择（红色）
        for i, opt in enumerate(opts):
            item = self.opt_buttons[i]
            if opt.get('correct'):
                item.config(bg="#1a4a3e", highlightbackground="#1a4a3e")
                if hasattr(item, 'letter_lbl'):
                    item.letter_lbl.config(bg="#4ecca3", fg="#1a1a2e")
                if hasattr(item, 'txt_lbl'):
                    item.txt_lbl.config(bg="#1a4a3e", fg="#4ecca3")
            elif i in selected_indices and not is_correct:
                item.config(bg="#4a1a2e", highlightbackground="#4a1a2e")
                if hasattr(item, 'letter_lbl'):
                    item.letter_lbl.config(bg="#e94560", fg="#fff")
                if hasattr(item, 'txt_lbl'):
                    item.txt_lbl.config(bg="#4a1a2e", fg="#e94560")
        if not is_correct:
            wrong_answers = [opts[i].get('text', '') for i in selected_indices]
            correct_answers = [o.get('text', '') for o in opts if o.get('correct')]
            if not any(w.get('id') == q.get('id') for w in self.wrong_notes):
                self.wrong_notes.append({'id': q.get('id'), 'text': q.get('text'),
                    'wrongAnswer': ' / '.join(wrong_answers),
                    'correctAnswer': ' / '.join(correct_answers),
                    'explanation': q.get('explanation', ''), 'category': q.get('category', '')})
        else:
            self.wrong_notes = [w for w in self.wrong_notes if w.get('id') != q.get('id')]
        if self.settings.get('showExp') and q.get('explanation'):
            self.exp_label.config(text=f"📖 {q.get('explanation', '')}")
            self.exp_label.pack(fill="x", padx=14, pady=(10, 0))
        self.submit_btn.pack_forget()
        # 答对且开启自动跳转时直接跳下一题
        if is_correct and self.settings.get('autoNext'):
            self.next_practice()
            return
        self.next_btn.pack()
        if self.role == "teacher":
            self._save_question_to_bank(q)
        else:
            self._save_student_progress()
        self.users['wrongNotes'] = self.wrong_notes
        save_users(self.users)
        self.update_top_stats()
    def _save_question_to_bank(self, q):
        for bname in [b['name'] for b in list_banks()]:
            bank = load_bank(bname)
            for i, bq in enumerate(bank):
                if bq.get('id') == q.get('id'):
                    bank[i] = q
                    save_bank(bname, bank)
                    return
    def _save_student_progress(self):
        """保存学员进度到服务器"""
        if not SERVER_URL or self.role != "student":
            return
        # 找出当前刷的是哪个班级题库
        if self.current_bank and str(self.current_bank).startswith("class_"):
            real_name = self.current_bank.replace("class_", "", 1)
            # 只保存有 done 状态的题目
            progress = [q for q in self.practice_queue if isinstance(q, dict)]
            server_request("POST", "/save_bank_questions", {
                "phone": self.phone, "bank_name": real_name, "questions": progress})
            server_request("POST", "/save_wrong_notes", {
                "phone": self.phone, "wrong_notes": self.wrong_notes})
    def edit_current_question(self):
        """练习时直接编辑当前题目"""
        if self.practice_idx >= len(self.practice_queue):
            return
        q = self.practice_queue[self.practice_idx]
        old_id = q.get('id', 0)
        old_text = q.get('text', '')
        
        # 找到这个题目属于哪个题库
        bank_name = self.current_bank
        if not bank_name:
            # 如果是从全部题库练习的，需要找到题目所属题库
            for b in list_banks():
                bank = load_bank(b['name'])
                for i, bq in enumerate(bank):
                    # 优先用 id 匹配，没 id 就用题目文本匹配
                    if old_id and bq.get('id') == old_id:
                        bank_name = b['name']
                        break
                    if not old_id and bq.get('text') == old_text:
                        bank_name = b['name']
                        break
                if bank_name:
                    break
        
        if not bank_name:
            messagebox.showerror("错误", "找不到题目所属题库")
            return
        
        bank = load_bank(bank_name)
        idx_in_bank = -1
        for i, bq in enumerate(bank):
            if old_id and bq.get('id') == old_id:
                idx_in_bank = i
                break
            if not old_id and bq.get('text') == old_text:
                idx_in_bank = i
                break
        
        if idx_in_bank == -1:
            messagebox.showerror("错误", "题目在题库中不存在")
            return
        
        # 打开编辑窗口
        win = tk.Toplevel(self.root)
        win.title(f"编辑第{idx_in_bank+1}题 - {bank_name}")
        win.geometry("520x560")
        win.configure(bg="#1a1a2e")
        win.resizable(False, False)
        win.transient(self.root)
        win.grab_set()
        
        qtype = q.get("type", "single")
        opts = q.get("options", [])
        
        tk.Label(win, text="题目类型", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(10, 2))
        type_frame = tk.Frame(win, bg="#1a1a2e")
        type_frame.pack()
        qtype_var = tk.StringVar(value=qtype)
        for val, txt in [("single", "单选题"), ("multiple", "多选题"), ("judge", "判断题")]:
            tk.Radiobutton(type_frame, text=txt, variable=qtype_var, value=val,
                          font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#eee",
                          selectcolor="#1a1a2e", activebackground="#1a1a2e").pack(side="left", padx=6)
        
        tk.Label(win, text="题目", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        q_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                         insertbackground="#eee", bd=0, highlightthickness=1,
                         highlightbackground="#0f3460", height=3, wrap="word")
        q_entry.insert("1.0", q.get("text", ""))
        q_entry.pack(fill="x", padx=30, ipady=4)
        
        tk.Label(win, text="选项（每行一个）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        opt_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                           insertbackground="#eee", bd=0, highlightthickness=1,
                           highlightbackground="#0f3460", height=5, wrap="word")
        if qtype == "judge":
            correct_ans = [chr(65+i) for i, o in enumerate(opts) if o.get('correct')]
            opt_txt = '正确' + ('[正确]' if 'A' in correct_ans else '') + '\n错误' + ('[正确]' if 'B' in correct_ans else '')
        else:
            opt_txt = '\n'.join([o.get('text', '') for o in opts])
        opt_entry.insert("1.0", opt_txt)
        opt_entry.pack(fill="x", padx=30, ipady=4)
        
        tk.Label(win, text="答案（多选用ABC，判断A正确B错误）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        correct = [chr(65+i) for i, o in enumerate(opts) if o.get("correct")]
        ans_entry = tk.Entry(win, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        ans_entry.insert(0, ",".join(correct))
        ans_entry.pack(fill="x", padx=30, ipady=5)
        
        tk.Label(win, text="解析", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        exp_entry = tk.Entry(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        exp_entry.insert(0, q.get("explanation", ""))
        exp_entry.pack(fill="x", padx=30, ipady=4)
        
        msg = tk.Label(win, text="", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#e94560")
        msg.pack(pady=(4, 2))
        
        def do_save():
            qtext = q_entry.get("1.0", "end-1c").strip()
            new_type = qtype_var.get()
            ans = ans_entry.get().strip()
            exp = exp_entry.get().strip()
            if not qtext:
                msg.config(text="请输入题目")
                return
            if not ans:
                msg.config(text="请输入答案")
                return
            
            bank = load_bank(bank_name)
            if new_type == "judge":
                new_opts = [{"text": "正确", "correct": ans.upper() == "A"},
                           {"text": "错误", "correct": ans.upper() == "B"}]
                bank[idx_in_bank] = {"text": qtext, "category": "判断题", "type": "judge",
                                    "options": new_opts, "explanation": exp,
                                    "done": False, "correct": False,
                                    "id": old_id or int(random.random() * 10**12)}
            else:
                opt_lines = [l.strip() for l in opt_entry.get("1.0", "end-1c").strip().split("\n") if l.strip()]
                if len(opt_lines) < 2:
                    msg.config(text="至少输入2个选项")
                    return
                new_opts = [{"text": o, "correct": False} for o in opt_lines]
                ans_upper = ans.upper().replace(",", "")
                for c in ans_upper:
                    if c.isalpha() and ord(c) - 65 < len(new_opts):
                        new_opts[ord(c) - 65]["correct"] = True
                if not any(o["correct"] for o in new_opts):
                    msg.config(text="答案不匹配")
                    return
                cat = "多选题" if new_type == "multiple" else "单选题"
                bank[idx_in_bank] = {"text": qtext, "category": cat, "type": new_type,
                                    "options": new_opts, "explanation": exp,
                                    "done": False, "correct": False,
                                    "id": old_id or int(random.random() * 10**12)}
            save_bank(bank_name, bank)
            
            # 更新当前练习队列中的题目
            self.practice_queue[self.practice_idx] = dict(bank[idx_in_bank])
            
            messagebox.showinfo("成功", "题目已更新！")
            win.destroy()
            # 刷新当前题目显示
            self.show_practice_question()
        
        tk.Button(win, text=" 保 存 ", font=("Microsoft YaHei", 12, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=20, pady=6,
                 cursor="hand2", relief="flat", command=do_save).pack(pady=(10, 4))
        
        def do_delete():
            if messagebox.askyesno("确认删除", f"确定删除此题吗？"):
                bank = load_bank(bank_name)
                bank.pop(idx_in_bank)
                save_bank(bank_name, bank)
                win.destroy()
                # 如果当前练习队列中有这个题，移除它
                if self.practice_idx < len(self.practice_queue):
                    q_in_queue = self.practice_queue[self.practice_idx]
                    queue_id = q_in_queue.get('id') or q_in_queue.get('text')
                    target_id = q.get('id') or q.get('text')
                    if queue_id == target_id:
                        self.practice_queue.pop(self.practice_idx)
                        if self.practice_idx >= len(self.practice_queue):
                            self.practice_idx = len(self.practice_queue) - 1
                self.update_top_stats()
                self.show_practice_question()
        
        tk.Button(win, text="删除此题", font=("Microsoft YaHei", 10),
                 bg="#e94560", fg="#fff", bd=0, padx=12, pady=4,
                 cursor="hand2", relief="flat", command=do_delete).pack()
    def next_practice(self):
        self.practice_idx += 1
        self.show_practice_question()
    def show_practice_complete(self):
        # 学员保存进度
        if self.role == "student":
            self._save_student_progress()
        self.clear_content()
        total = len(self.practice_queue)
        correct = len([q for q in self.practice_queue if q.get('correct')])
        acc = round(correct / total * 100) if total > 0 else 0
        emoji = "🎉" if acc >= 80 else "👍" if acc >= 60 else "💪"
        title = "太棒了！" if acc >= 80 else "还不错！" if acc >= 60 else "继续加油！"
        tk.Label(self.content_frame, text=emoji, font=("Microsoft YaHei", 56), bg="#1a1a2e", fg="#eee").pack(pady=(60, 10))
        tk.Label(self.content_frame, text=title, font=("Microsoft YaHei", 22, "bold"), bg="#1a1a2e", fg="#eee").pack()
        tk.Label(self.content_frame, text=f"完成 {total} 题，正确率 {acc}%（{correct}/{total}）", font=("Microsoft YaHei", 12), bg="#1a1a2e", fg="#888").pack(pady=10)
        btn_row = tk.Frame(self.content_frame, bg="#1a1a2e")
        btn_row.pack(pady=20)
        ttk.Button(btn_row, text="🔄 再刷一遍", style="Primary.TButton", command=self.show_type_practice).pack(side="left", padx=6)
        ttk.Button(btn_row, text="🏠 返回首页", style="Secondary.TButton", command=lambda: self.navigate("home")).pack(side="left", padx=6)
    # ===== 错题本 =====
    def show_wrong(self):
        self.clear_content()
        tk.Label(self.content_frame, text="📝 错题本", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 14))
        if not self.wrong_notes:
            tk.Label(self.content_frame, text="🎉\n\n太棒了！暂无错题记录！", font=("Microsoft YaHei", 14), bg="#1a1a2e", fg="#888", justify="center").pack(pady=60)
            ttk.Button(self.content_frame, text="去练习", style="Secondary.TButton", command=lambda: self.navigate("type_practice")).pack()
            return
        canvas = tk.Canvas(self.content_frame, bg="#1a1a2e", bd=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content_frame, orient="vertical", command=canvas.yview)
        sf = tk.Frame(canvas, bg="#1a1a2e")
        sf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=sf, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        for i, w in enumerate(self.wrong_notes):
            item = tk.Frame(sf, bg="#16213e")
            item.pack(fill="x", pady=4, padx=2)
            tk.Frame(item, bg="#e94560", width=4).pack(side="left", fill="y")
            inner = tk.Frame(item, bg="#16213e")
            inner.pack(side="left", fill="both", expand=True, padx=10, pady=8)
            tk.Label(inner, text=f"{i+1}. {w.get('text', '')[:80]}{'...' if len(w.get('text', '')) > 80 else ''}",
                    font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee", anchor="w", wraplength=600).pack(anchor="w")
            tk.Label(inner, text=f"✗ {w.get('wrongAnswer', '')}", font=("Microsoft YaHei", 10), bg="#16213e", fg="#e94560", anchor="w").pack(anchor="w", pady=(4, 1))
            tk.Label(inner, text=f"✓ {w.get('correctAnswer', '')}", font=("Microsoft YaHei", 10), bg="#16213e", fg="#4ecca3", anchor="w").pack(anchor="w")
            if w.get('explanation'):
                tk.Label(inner, text=f"📖 {w.get('explanation', '')}", font=("Microsoft YaHei", 9), bg="#16213e", fg="#888", anchor="w", wraplength=600).pack(anchor="w", pady=(3, 0))
            tk.Button(inner, text="✕ 移除", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#888", bd=0, padx=8, cursor="hand2",
                     command=lambda wid=w.get('id'): (setattr(self, 'wrong_notes', [x for x in self.wrong_notes if x.get('id') != wid]),
                     setattr(self.users, 'wrongNotes', self.wrong_notes), save_users(self.users), self.show_wrong())).pack(anchor="e", pady=(6, 0))
    # ===== 题库管理 =====
    def show_banks(self):
        self.clear_content()
        banks = list_banks()
        tk.Label(self.content_frame, text="📚 题库管理", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 14))
        if not banks:
            tk.Label(self.content_frame, text="暂无题库", font=("Microsoft YaHei", 12), bg="#1a1a2e", fg="#888").pack(pady=30)
            ttk.Button(self.content_frame, text="导入题库", style="Primary.TButton", command=lambda: self.navigate("import")).pack()
            return
        canvas = tk.Canvas(self.content_frame, bg="#1a1a2e", bd=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content_frame, orient="vertical", command=canvas.yview)
        sf = tk.Frame(canvas, bg="#1a1a2e")
        sf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=sf, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        for b in banks:
            item = tk.Frame(sf, bg="#16213e")
            item.pack(fill="x", pady=4, padx=2)
            color = b['types'].get('judge') and "#67c23a" or b['types'].get('multiple') and "#e6a23c" or "#4a90d9"
            tk.Frame(item, bg=color, width=4).pack(side="left", fill="y")
            inner = tk.Frame(item, bg="#16213e")
            inner.pack(side="left", fill="both", expand=True, padx=12, pady=10)
            names = {'single': '单选', 'multiple': '多选', 'judge': '判断', 'fill': '填空', 'essay': '问答'}
            tk.Label(inner, text=f"📚 {b['name']}", font=("Microsoft YaHei", 12, "bold"), bg="#16213e", fg="#eee").pack(anchor="w")
            tk.Label(inner, text=" | ".join([f"{names.get(t,t)} {c}题" for t, c in b['types'].items()]), font=("Microsoft YaHei", 9), bg="#16213e", fg="#888").pack(anchor="w", pady=(2, 6))
            prog_bg = tk.Frame(inner, bg="#0f3460", height=6)
            prog_bg.pack(fill="x", pady=(0, 6))
            prog_bg.pack_propagate(False)
            if b['done'] > 0:
                tk.Frame(prog_bg, bg="#e94560", height=6).place(x=0, y=0, relwidth=b['done'] / b['total'], relheight=1)
            tk.Label(inner, text=f"进度: {b['done']}/{b['total']} | 正确率: {b['acc']}%", font=("Microsoft YaHei", 9), bg="#16213e", fg="#666").pack(anchor="w")
            btn_row = tk.Frame(inner, bg="#16213e")
            btn_row.pack(anchor="e", pady=(6, 0))
            for lbl, cmd in [("练习", lambda n=b['name']: (setattr(self, 'current_bank', n), setattr(self, 'current_type_filter', None), self.navigate("type_practice"))),
                             ("重置", lambda n=b['name']: (setattr(self, 'current_bank', n), self.reset_bank_progress(n)))]:
                tk.Button(btn_row, text=f"✍️ {lbl}" if lbl == "练习" else f"🔄 {lbl}", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#aaa", bd=0, padx=8, pady=3, cursor="hand2", relief="flat", command=cmd).pack(side="left", padx=2)
            tk.Button(btn_row, text="🗑️", font=("Microsoft YaHei", 9), bg="#16213e", fg="#e94560", bd=0, padx=8, cursor="hand2", relief="flat",
                     command=lambda n=b['name']: (delete_bank(n), self.show_banks(), self.update_top_stats())).pack(side="left", padx=2)
            tk.Button(btn_row, text=" 管理题目", font=("Microsoft YaHei", 9), bg="#4ecca3", fg="#1a1a2e", bd=0, padx=8, pady=2,
                     cursor="hand2", relief="flat",
                     command=lambda n=b['name']: self.show_bank_questions(n)).pack(side="left", padx=2)
    def reset_bank_progress(self, bank_name):
        if messagebox.askyesno("重置进度", f"确定要重置「{bank_name}」的做题进度吗？"):
            bank = load_bank(bank_name)
            for q in bank:
                q['done'] = False
                q['correct'] = False
            save_bank(bank_name, bank)
            self.update_top_stats()
            messagebox.showinfo("完成", "进度已重置")
            self.show_banks()
    # ===== 统计 =====
    def show_stats(self):
        self.clear_content()
        banks = list_banks()
        total = sum(b['total'] for b in banks)
        done = sum(b['done'] for b in banks)
        correct = sum(b['correct'] for b in banks)
        acc = round(correct / done * 100) if done > 0 else 0
        tk.Label(self.content_frame, text="📊 学习统计", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 14))
        grid = tk.Frame(self.content_frame, bg="#1a1a2e")
        grid.pack(fill="x", pady=(0, 14))
        for val, lbl, color in [(len(banks), "题库数", "#e94560"), (total, "总题数", "#e94560"), (done, "已完成", "#4ecca3"), (f"{acc}%", "正确率", "#f9c74f")]:
            c = tk.Frame(grid, bg="#16213e")
            c.pack(side="left", padx=6, fill="both", expand=True)
            tk.Label(c, text=str(val), font=("Microsoft YaHei", 26, "bold"), bg="#16213e", fg=color).pack(pady=(14, 4))
            tk.Label(c, text=lbl, font=("Microsoft YaHei", 11), bg="#16213e", fg="#888").pack(pady=(0, 14))
        if banks:
            tk.Label(self.content_frame, text="各题库统计", font=("Microsoft YaHei", 12), bg="#1a1a2e", fg="#888").pack(anchor="w", pady=(6, 8))
            c = self.card()
            for b in banks:
                row = tk.Frame(c, bg="#16213e")
                row.pack(fill="x", pady=4, padx=10)
                color = b['types'].get('judge') and "#67c23a" or b['types'].get('multiple') and "#e6a23c" or "#4a90d9"
                tk.Label(row, text=f"📚 {b['name']}", font=("Microsoft YaHei", 10), bg="#16213e", fg=color, width=14, anchor="w").pack(side="left")
                tk.Label(row, text=f"{b['done']}/{b['total']}题 · {b['acc']}%正确率", font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(side="left")
                prog_bg = tk.Frame(row, bg="#0f3460", width=100, height=6)
                prog_bg.pack(side="left", pady=4, padx=10)
                prog_bg.pack_propagate(False)
                if b['done'] > 0:
                    tk.Frame(prog_bg, bg="#e94560", height=6).place(x=0, y=0, relwidth=b['done'] / b['total'], relheight=1)
        btn_row = tk.Frame(self.content_frame, bg="#1a1a2e")
        btn_row.pack(pady=14)
        ttk.Button(btn_row, text="🔄 重置全部进度", style="Secondary.TButton", command=self.reset_progress).pack(side="left", padx=4)
        ttk.Button(btn_row, text="🗑️ 清空全部数据", style="Secondary.TButton", command=self.clear_all_data).pack(side="left", padx=4)
    def reset_progress(self):
        if messagebox.askyesno("确认", "确定要重置所有做题进度吗？题库会保留。"):
            for bname in [b['name'] for b in list_banks()]:
                bank = load_bank(bname)
                for q in bank:
                    q['done'] = False
                    q['correct'] = False
                save_bank(bname, bank)
            self.wrong_notes = []
            self.users['wrongNotes'] = []
            save_users(self.users)
            self.update_top_stats()
            messagebox.showinfo("完成", "进度已重置")
            self.show_stats()
    def clear_all_data(self):
        if messagebox.askyesno("危险操作", "⚠️ 确定要清空所有数据吗？题库和错题记录将全部删除，此操作不可恢复！"):
            for bname in [b['name'] for b in list_banks()]:
                delete_bank(bname)
            self.wrong_notes = []
            self.users['wrongNotes'] = []
            save_users(self.users)
            self.update_top_stats()
            messagebox.showinfo("完成", "所有数据已清空")
            self.navigate("home")
    # ===== 设置 =====
    def show_settings(self):
        self.clear_content()
        tk.Label(self.content_frame, text="⚙️ 设置", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 14))
        c = self.card()
        row = tk.Frame(c, bg="#16213e")
        row.pack(fill="x", padx=14, pady=10)
        tk.Label(row, text="随机出题", font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee").pack(side="left")
        self.shuffle_var = tk.BooleanVar(value=self.settings.get("shuffle", False))
        tk.Checkbutton(row, variable=self.shuffle_var, bg="#16213e", activebackground="#16213e", selectcolor="#4ecca3", bd=0).pack(side="right")
        row2 = tk.Frame(c, bg="#16213e")
        row2.pack(fill="x", padx=14, pady=10)
        tk.Label(row2, text="答题后显示解析", font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee").pack(side="left")
        self.showexp_var = tk.BooleanVar(value=self.settings.get("showExp", True))
        tk.Checkbutton(row2, variable=self.showexp_var, bg="#16213e", activebackground="#16213e", selectcolor="#4ecca3", bd=0).pack(side="right")
        row3 = tk.Frame(c, bg="#16213e")
        row3.pack(fill="x", padx=14, pady=10)
        tk.Label(row3, text="答对自动跳转下一题", font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee").pack(side="left")
        self.auto_next_var = tk.BooleanVar(value=self.settings.get("autoNext", False))
        tk.Checkbutton(row3, variable=self.auto_next_var, bg="#16213e", activebackground="#16213e", selectcolor="#4ecca3", bd=0).pack(side="right")
        tk.Button(c, text="✓ 保存设置", font=("Microsoft YaHei", 11, "bold"), bg="#e94560", fg="#fff", bd=0, padx=16, pady=7, cursor="hand2", relief="flat", command=self.save_settings).pack(pady=(6, 0), anchor="e")
        c2 = self.card()
        tk.Label(c2, text="数据管理", font=("Microsoft YaHei", 12), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(10, 8))
        btn_row = tk.Frame(c2, bg="#16213e")
        btn_row.pack(fill="x", padx=12, pady=(0, 10))
        ttk.Button(btn_row, text="📤 导出题库", style="Secondary.TButton", command=self.do_export).pack(side="left", padx=4)
        ttk.Button(btn_row, text="📥 导入题库", style="Secondary.TButton", command=lambda: self.navigate("import")).pack(side="left", padx=4)
    def save_settings(self):
        self.settings["shuffle"] = self.shuffle_var.get()
        self.settings["showExp"] = self.showexp_var.get()
        self.settings["autoNext"] = self.auto_next_var.get()
        self.users['settings'] = self.settings
        save_users(self.users)
        messagebox.showinfo("成功", "设置已保存！")
    # ===== 导入 =====
    def show_import(self):
        self.clear_content()
        tk.Label(self.content_frame, text="📥 导入题库", font=("Microsoft YaHei", 18, "bold"), bg="#1a1a2e", fg="#eee").pack(anchor="w", pady=(0, 14))
        c = self.card()
        tk.Label(c, text="支持格式：Excel(.xlsx) / Word(.docx) / 文本(.txt)", font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(anchor="w", padx=14, pady=(10, 6))
        btn_row = tk.Frame(c, bg="#16213e")
        btn_row.pack(fill="x", padx=12, pady=10)
        ttk.Button(btn_row, text="📂 从文档导入", style="Success.TButton", command=self.do_import_file).pack(side="left", padx=4)
        ttk.Button(btn_row, text="📥 JSON导入", style="Primary.TButton", command=self.show_json_import).pack(side="left", padx=4)
        ttk.Button(btn_row, text="📤 导出题库", style="Secondary.TButton", command=self.do_export).pack(side="left", padx=4)
        doc_info = tk.Frame(c, bg="#16213e")
        doc_info.pack(fill="x", padx=12, pady=(0, 10))
        tk.Label(doc_info, text="📂 文档导入支持格式：", font=("Microsoft YaHei", 10), bg="#16213e", fg="#888").pack(anchor="w")
        for fmt in ["Excel - 第1行标题，第2行表头，第3行起数据（答案/题目/选项列）",
                    "Word - 题目以数字编号，选项A.B.C.D.，答案标注在选项",
                    "TXT - 每题以题号开头，选项A.B.C.D.，答案：xxx"]:
            tk.Label(doc_info, text=f"  • {fmt}", font=("Microsoft YaHei", 9), bg="#16213e", fg="#666", anchor="w", justify="left").pack(anchor="w")
        self.import_result = tk.Label(c, text="", font=("Microsoft YaHei", 10), bg="#16213e", fg="#4ecca3")
        self.import_result.pack(anchor="w", padx=14, pady=(0, 8))
    def show_json_import(self):
        win = tk.Toplevel(self.root)
        win.title("JSON 导入")
        win.geometry("600x500")
        win.configure(bg="#1a1a2e")
        win.transient(self.root)
        win.grab_set()
        tk.Label(win, text="粘贴 JSON 数据（数组格式，每项包含 text, options, category 等）", font=("Microsoft YaHei", 10), bg="#1a1a2e", fg="#888").pack(anchor="w", padx=16, pady=10)
        text = tk.Text(win, font=("Consolas", 10), bg="#0f3460", fg="#aaddaa", bd=0, insertbackground="#fff")
        text.pack(fill="both", expand=True, padx=12, pady=4)
        sample = json.dumps([{"text": "示例题目", "category": "分类", "type": "single",
                             "options": [{"text": "选项A", "correct": True}, {"text": "选项B", "correct": False}]}], ensure_ascii=False, indent=2)
        tk.Button(win, text="📋 填写示例", font=("Microsoft YaHei", 9), bg="#0f3460", fg="#888", bd=0, padx=10, cursor="hand2",
                 command=lambda: text.insert("1.0", sample)).pack(anchor="w", padx=12)
        result = [None]
        def do_import():
            bank_name = self._ask_bank_name()
            if not bank_name:
                return
            try:
                data = json.loads(text.get("1.0", "end").strip())
                if not isinstance(data, list):
                    raise ValueError("需要数组格式")
                cnt = 0
                for item in data:
                    if item.get("text") and item.get("options"):
                        item['id'] = int(random.random() * 10**12)
                        item['done'] = False
                        item['correct'] = False
                        cnt += 1
                bank = load_bank(bank_name)
                bank.extend(data)
                save_bank(bank_name, bank)
                self.update_top_stats()
                result[0] = f"✅ 成功导入 {cnt} 道题目到「{bank_name}」！"
                win.destroy()
            except Exception as e:
                result[0] = f"❌ 失败：{e}"
        btn_row = tk.Frame(win, bg="#1a1a2e")
        btn_row.pack(pady=10)
        tk.Button(btn_row, text="确认导入", font=("Microsoft YaHei", 11, "bold"), bg="#e94560", fg="#fff", bd=0, padx=20, pady=6, cursor="hand2", command=do_import).pack(side="left", padx=8)
        tk.Button(btn_row, text="取消", font=("Microsoft YaHei", 11), bg="#0f3460", fg="#aaa", bd=0, padx=20, pady=6, cursor="hand2", command=win.destroy).pack(side="left", padx=8)
        win.wait_window()
        if result[0]:
            self.import_result.config(text=result[0], fg="#4ecca3" if "✅" in result[0] else "#e94560")
    def do_import_file(self):
        path = filedialog.askopenfilename(title="选择文档导入", filetypes=[("支持格式", "*.docx;*.xlsx;*.xls;*.txt"), ("Word", "*.docx"), ("Excel", "*.xlsx;*.xls"), ("TXT", "*.txt"), ("所有文件", "*.*")])
        if not path:
            return
        try:
            cnt = self.import_document(path)
            self.import_result.config(text=f"✅ 从文档成功导入 {cnt} 道题目！", fg="#4ecca3")
        except Exception as e:
            self.import_result.config(text=str(e), fg="#e94560")
    def import_document(self, path):
        ext = os.path.splitext(path)[1].lower()
        bank_name = self._ask_bank_name()
        if not bank_name:
            raise Exception("已取消")
        try:
            if ext == '.docx':
                if not DOCX_AVAILABLE:
                    raise Exception("python-docx 未安装")
                items = self.parse_docx(path)
            elif ext in ['.xlsx', '.xls']:
                if not XLSX_AVAILABLE:
                    raise Exception("openpyxl 未安装")
                items = self.parse_xlsx(path)
            elif ext == '.txt':
                items = self.parse_txt(path)
            else:
                raise Exception(f"不支持的格式：{ext}")
            cnt = 0
            for item in items:
                if item.get('text') and item.get('options'):
                    item['id'] = int(random.random() * 10**12)
                    item['done'] = False
                    item['correct'] = False
                    cnt += 1
            bank = load_bank(bank_name)
            bank.extend(items)
            save_bank(bank_name, bank)
            self.update_top_stats()
            return cnt
        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            raise Exception(f"导入失败：[{type(e).__name__}] {e}\n{tb[:300]}")
    def _ask_bank_name(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("题库命名")
        dialog.geometry("400x160")
        dialog.configure(bg="#1a1a2e")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)
        result = [None]
        tk.Label(dialog, text="请输入题库名称：", font=("Microsoft YaHei", 12), bg="#1a1a2e", fg="#eee").pack(pady=(16, 8))
        entry = tk.Entry(dialog, font=("Microsoft YaHei", 12), bg="#0f3460", fg="#eee", bd=0, insertbackground="#fff", justify="center")
        entry.pack(fill="x", padx=30, pady=4)
        entry.focus()
        btn_frame = tk.Frame(dialog, bg="#1a1a2e")
        btn_frame.pack(pady=14)
        def confirm():
            name = entry.get().strip()
            if name:
                result[0] = name
                dialog.destroy()
        def cancel():
            dialog.destroy()
        tk.Button(btn_frame, text="确认", font=("Microsoft YaHei", 11, "bold"), bg="#e94560", fg="#fff", bd=0, padx=20, pady=6, cursor="hand2", command=confirm).pack(side="left", padx=8)
        tk.Button(btn_frame, text="取消", font=("Microsoft YaHei", 11), bg="#0f3460", fg="#aaa", bd=0, padx=20, pady=6, cursor="hand2", command=cancel).pack(side="left", padx=8)
        dialog.bind('<Return>', lambda e: confirm())
        dialog.bind('<Escape>', lambda e: cancel())
        dialog.wait_window()
        return result[0]
    def do_export(self):
        banks = list_banks()
        if not banks:
            messagebox.showwarning("提示", "暂无题库可导出")
            return
        path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON文件", "*.json")], initialfile="题库导出.json")
        if not path:
            return
        all_data = []
        for bname in [b['name'] for b in banks]:
            for q in load_bank(bname):
                d = dict(q)
                d['bank'] = bname
                all_data.append(d)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(all_data, f, ensure_ascii=False, indent=2)
        messagebox.showinfo("导出成功", f"共导出 {len(all_data)} 道题目到：\n{path}")
    # ===== 文档解析 ======
    def parse_docx(self, path):
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return self._parse_text_format('\n'.join(full_text))
    def parse_xlsx(self, path):
        questions = []
        wb = openpyxl.load_workbook(path, data_only=True)
        for sheet in wb.worksheets:
            sheet_name = sheet.title
            rows = list(sheet.iter_rows(values_only=True))
            for row in rows[2:]:
                if not row or not any(row):
                    continue
                row = list(row)
                while len(row) < 12:
                    row.append(None)
                # Excel: A(0)=题干, B(1)=题型, C-J(2-9)=选项A-H, K(10)=正确答案, L(11)=解析
                question = str(row[0]).strip() if row[0] else ''
                qtype_str = str(row[1]).strip() if row[1] else ''
                answer = str(row[10]).strip() if row[10] else ''
                explanation = str(row[11]).strip() if row[11] else ''
                if not question or question == 'None':
                    continue
                if '判断' in sheet_name or '判断' in qtype_str:
                    opts = [{'text': '正确', 'correct': answer.upper() == 'A'},
                            {'text': '错误', 'correct': answer.upper() == 'B'}]
                    questions.append({'text': question, 'category': '判断题', 'type': 'judge', 'options': opts, 'explanation': explanation})
                else:
                    opts = []
                    for col_idx in range(2, 10):
                        opt_text = str(row[col_idx]).strip() if col_idx < len(row) and row[col_idx] else ''
                        if opt_text and opt_text not in ['None', '']:
                            opts.append({'text': opt_text, 'correct': False})
                    if not opts:
                        continue
                    answer_upper = answer.upper()
                    for c in answer_upper:
                        if c.isalpha() and ord(c) - 65 < len(opts):
                            opts[ord(c) - 65]['correct'] = True
                    is_multi = '多选' in sheet_name or '多选' in qtype_str
                    qtype_name = '多选题' if is_multi else '单选题'
                    qtype_val = 'multiple' if is_multi else 'single'
                    questions.append({'text': question, 'category': qtype_name,
                                     'type': qtype_val, 'options': opts, 'explanation': explanation})
        return questions
    def parse_txt(self, path):
        for enc in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin1']:
            try:
                with open(path, 'r', encoding=enc, errors='ignore') as f:
                    content = f.read()
                if content and not content.startswith('<'):
                    return self._parse_text_format(content)
            except:
                continue
        with open(path, 'r', encoding='gbk', errors='ignore') as f:
            return self._parse_text_format(f.read())
    def _parse_text_format(self, content):
        questions = []
        lines = content.split('\n')
        current_section = '选择题'
        current_q = None
        current_opts = []
        current_exp = ""
        current_type = 'single'
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.lower().startswith('ps') or line.startswith('注：') or line.startswith('注意：'):
                continue
            is_section = False
            for key in ['一', '二', '三', '四', '五', '六']:
                if line.startswith(key + '.') or line.startswith(key + '、'):
                    if current_q:
                        self._add_question(questions, current_q, current_opts, current_section, current_type, current_exp)
                    current_q = None
                    current_opts = []
                    current_exp = ""
                    current_type = 'single'
                    if key == '一': current_section = '选择题'
                    elif key == '二': current_section = '判断题'; current_type = 'judge'
                    elif key == '三': current_section = '填空题'; current_type = 'fill'
                    elif key == '四': current_section = '形式与策略'; current_type = 'essay'
                    elif key == '五': current_section = '材料分析'; current_type = 'essay'
                    elif key == '六': current_section = '简答题'; current_type = 'essay'
                    is_section = True
                    break
            if is_section:
                continue
            if len(line) > 2 and line[0].isdigit() and line[1] == '.':
                if current_q:
                    self._add_question(questions, current_q, current_opts, current_section, current_type, current_exp)
                current_q = line[2:].strip()
                current_opts = []
                current_exp = ""
                current_type = 'single' if current_section == '选择题' else 'judge' if current_section == '判断题' else 'fill' if current_section == '填空题' else 'essay'
            elif current_q is not None:
                opt_markers = ['A.', 'B.', 'C.', 'D.', 'E.', 'a.', 'b.', 'c.', 'd.', 'A、', 'B、', 'C、', 'D、']
                found = next((m for m in opt_markers if line.startswith(m)), None)
                if found:
                    opt_text = line[len(found):].strip().replace('*', '').replace('√', '').replace('✓', '').strip()
                    current_opts.append({'text': opt_text, 'correct': '*' in line or '√' in line or '✓' in line})
                    continue
                if '答案' in line:
                    ans = line
                    for prefix in ['答案：', '答案:', '答：', '答:', '【答案】', '[答案]']:
                        if prefix in line:
                            ans = line.split(prefix, 1)[-1].strip()
                            break
                    if not ans:
                        continue
                    current_exp = ans
                    if current_section == '判断题':
                        correct_is_a = any(k in ans for k in ['正确', '对', 'T', '√', 'true'])
                        current_opts = [{'text': '正确', 'correct': correct_is_a}, {'text': '错误', 'correct': not correct_is_a}]
                    elif current_section == '填空题':
                        current_opts = [{'text': ans, 'correct': True}, {'text': '其他答案', 'correct': False}]
                    else:
                        for o in current_opts:
                            o['correct'] = False
                        for c in ans.upper():
                            if c.isalpha() and ord(c) - 65 < len(current_opts):
                                current_opts[ord(c) - 65]['correct'] = True
                        if not any(o['correct'] for o in current_opts):
                            for j in range(min(len(ans), len(current_opts))):
                                if ans[j].isalpha():
                                    current_opts[j]['correct'] = True
        if current_q:
            self._add_question(questions, current_q, current_opts, current_section, current_type, current_exp)
        return questions
    def _add_question(self, questions, text, opts, cat, qtype, exp):
        if not text or not opts:
            return
        if cat == '判断题' and (not opts or len(opts) < 2):
            correct_is_first = not any(k in exp for k in ['错误', '错', 'B', '×', 'F'])
            opts = [{'text': '正确', 'correct': correct_is_first}, {'text': '错误', 'correct': not correct_is_first}]
        if cat == '填空题' and (not opts or len(opts) < 2):
            opts = [{'text': exp or '有答案', 'correct': True}, {'text': '其他答案', 'correct': False}]
        if not opts:
            return
        if not any(o['correct'] for o in opts):
            opts[0]['correct'] = True
        questions.append({'text': text, 'category': cat, 'type': qtype, 'options': opts, 'explanation': exp})
    def show_add_question(self, bank_name):
        win = tk.Toplevel(self.root)
        win.title(f"添加题目 - {bank_name}")
        win.geometry("520x520")
        win.configure(bg="#1a1a2e")
        win.resizable(False, False)
        win.transient(self.root)
        win.grab_set()
        tk.Label(win, text="题目类型", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(12, 2))
        type_frame = tk.Frame(win, bg="#1a1a2e")
        type_frame.pack()
        qtype_var = tk.StringVar(value="single")
        for val, txt in [("single", "单选题"), ("multiple", "多选题"), ("judge", "判断题")]:
            tk.Radiobutton(type_frame, text=txt, variable=qtype_var, value=val,
                          font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#eee",
                          selectcolor="#1a1a2e", activebackground="#1a1a2e").pack(side="left", padx=6)
        tk.Label(win, text="题目内容", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(8, 2))
        q_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                         insertbackground="#eee", bd=0, highlightthickness=1,
                         highlightbackground="#0f3460", height=3, wrap="word")
        q_entry.pack(fill="x", padx=30, ipady=4)
        tk.Label(win, text="选项（每行一个）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(8, 2))
        opt_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                           insertbackground="#eee", bd=0, highlightthickness=1,
                           highlightbackground="#0f3460", height=5, wrap="word")
        opt_entry.pack(fill="x", padx=30, ipady=4)
        tk.Label(win, text="答案（多选用ABC格式，判断题A=正确B=错误）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(8, 2))
        ans_entry = tk.Entry(win, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        ans_entry.pack(fill="x", padx=30, ipady=5)
        tk.Label(win, text="解析（选填）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(8, 2))
        exp_entry = tk.Entry(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        exp_entry.pack(fill="x", padx=30, ipady=4)
        msg = tk.Label(win, text="", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#e94560")
        msg.pack(pady=(6, 2))
        def do_save():
            qtext = q_entry.get("1.0", "end-1c").strip()
            qtype = qtype_var.get()
            ans = ans_entry.get().strip()
            exp = exp_entry.get().strip()
            if not qtext:
                msg.config(text="请输入题目")
                return
            if not ans:
                msg.config(text="请输入答案")
                return
            bank = load_bank(bank_name)
            if qtype == "judge":
                opts = [{"text": "正确", "correct": ans.upper() == "A"},
                        {"text": "错误", "correct": ans.upper() == "B"}]
                self._add_question(bank, qtext, opts, "判断题", "judge", exp)
            else:
                opt_lines = [l.strip() for l in opt_entry.get("1.0", "end-1c").strip().split("\n") if l.strip()]
                if len(opt_lines) < 2:
                    msg.config(text="至少输入2个选项，每行一个")
                    return
                opts = [{"text": o, "correct": False} for o in opt_lines]
                ans_upper = ans.upper()
                for c in ans_upper:
                    if c.isalpha() and ord(c) - 65 < len(opts):
                        opts[ord(c) - 65]["correct"] = True
                if not any(o["correct"] for o in opts):
                    msg.config(text="答案字母与选项不匹配")
                    return
                cat_name = "多选题" if qtype == "multiple" else "单选题"
                self._add_question(bank, qtext, opts, cat_name, qtype, exp)
            save_bank(bank_name, bank)
            messagebox.showinfo("成功", "题目已添加到题库")
            win.destroy()
            self.update_top_stats()
        tk.Button(win, text=" 保 存 ", font=("Microsoft YaHei", 12, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=20, pady=6,
                 cursor="hand2", relief="flat", command=do_save).pack(pady=(12, 6))
    def show_bank_questions(self, bank_name):
        """显示题库的所有题目列表，支持点击编辑和自由练习"""
        self.clear_content()
        self.current_bank = bank_name
        
        bank = load_bank(bank_name)
        total = len(bank)
        
        # 顶部
        header = tk.Frame(self.content_frame, bg="#1a1a2e")
        header.pack(fill="x", pady=(4, 8))
        
        row0 = tk.Frame(header, bg="#1a1a2e")
        row0.pack(fill="x")
        tk.Label(row0, text=f"题库: {bank_name} ({total}题)",
                font=("Microsoft YaHei", 15, "bold"), bg="#1a1a2e", fg="#eee").pack(side="left", padx=8)
        tk.Button(row0, text="< 返回", font=("Microsoft YaHei", 9),
                 bg="#16213e", fg="#888", bd=0, padx=8, pady=2,
                 cursor="hand2", relief="flat",
                 command=self.show_banks).pack(side="right", padx=8)
        
        # 操作按钮行
        act_row = tk.Frame(header, bg="#1a1a2e")
        act_row.pack(fill="x", pady=(4, 0))
        # 自由练习 - 输入起始题号
        tk.Label(act_row, text="从第", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#aaa").pack(side="left", padx=(8, 2))
        start_var = tk.StringVar(value="1")
        sp = tk.Spinbox(act_row, from_=1, to=total, textvariable=start_var,
                       width=4, font=("Microsoft YaHei", 9),
                       bg="#16213e", fg="#eee", bd=0, buttonbackground="#0f3460",
                       readonlybackground="#16213e")
        sp.pack(side="left", padx=2)
        tk.Label(act_row, text="题开始练习", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#aaa").pack(side="left", padx=2)
        tk.Button(act_row, text=" 开始练习 ", font=("Microsoft YaHei", 9, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=12, pady=2,
                 cursor="hand2", relief="flat",
                 command=lambda: self.start_bank_at(bank_name, int(start_var.get()) - 1)).pack(side="left", padx=8)
        tk.Button(act_row, text=" + 添加题目 ", font=("Microsoft YaHei", 9),
                 bg="#4a90d9", fg="#fff", bd=0, padx=12, pady=2,
                 cursor="hand2", relief="flat",
                 command=lambda: self.show_add_question(bank_name)).pack(side="right", padx=8)
        
        # 可滚动题目列表
        canvas = tk.Canvas(self.content_frame, bg="#1a1a2e", bd=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self.content_frame, orient="vertical", command=canvas.yview)
        sf = tk.Frame(canvas, bg="#1a1a2e")
        sf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=sf, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        for idx, q in enumerate(bank):
            if not isinstance(q, dict):
                continue
            item = tk.Frame(sf, bg="#16213e")
            item.pack(fill="x", pady=3, padx=4)
            
            # 类型色条
            type_colors = {"single": "#4a90d9", "multiple": "#e6a23c", "judge": "#67c23a"}
            tc = type_colors.get(q.get("type", ""), "#888")
            tk.Frame(item, bg=tc, width=4).pack(side="left", fill="y")
            
            inner = tk.Frame(item, bg="#16213e")
            inner.pack(side="left", fill="both", expand=True, padx=8, pady=6)
            
            # 题号 + 题干
            qtext = q.get("text", "")[:60]
            if len(q.get("text", "")) > 60:
                qtext += "..."
            tk.Label(inner, text=f"{idx+1}. {qtext}",
                    font=("Microsoft YaHei", 10), bg="#16213e", fg="#eee",
                    anchor="w", justify="left").pack(anchor="w", fill="x")
            
            # 答案摘要
            opts = q.get("options", [])
            correct_ans = [chr(65+i) for i, o in enumerate(opts) if o.get("correct")]
            ans_text = ",".join(correct_ans) if correct_ans else "-"
            type_names = {"single": "单选", "multiple": "多选", "judge": "判断"}
            tname = type_names.get(q.get("type", ""), q.get("type", ""))
            tk.Label(inner, text=f"[{tname}] 答案: {ans_text}",
                    font=("Microsoft YaHei", 8), bg="#16213e", fg="#888").pack(anchor="w")
            
            # 编辑按钮
            tk.Button(inner, text="编辑", font=("Microsoft YaHei", 9),
                     bg="#0f3460", fg="#4ecca3", bd=0, padx=10, pady=2,
                     cursor="hand2", relief="flat",
                     command=lambda q=q, idx=idx: self.edit_question(bank_name, q, idx)).pack(anchor="e", pady=(4, 0))
    
    def edit_question(self, bank_name, question, idx):
        """编辑单道题目"""
        win = tk.Toplevel(self.root)
        win.title(f"编辑第{idx+1}题 - {bank_name}")
        win.geometry("520x560")
        win.configure(bg="#1a1a2e")
        win.resizable(False, False)
        win.transient(self.root)
        win.grab_set()
        
        qtype = question.get("type", "single")
        opts = question.get("options", [])
        
        tk.Label(win, text="题目类型", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(10, 2))
        type_frame = tk.Frame(win, bg="#1a1a2e")
        type_frame.pack()
        qtype_var = tk.StringVar(value=qtype)
        for val, txt in [("single", "单选题"), ("multiple", "多选题"), ("judge", "判断题")]:
            tk.Radiobutton(type_frame, text=txt, variable=qtype_var, value=val,
                          font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#eee",
                          selectcolor="#1a1a2e", activebackground="#1a1a2e").pack(side="left", padx=6)
        
        tk.Label(win, text="题目", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        q_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                         insertbackground="#eee", bd=0, highlightthickness=1,
                         highlightbackground="#0f3460", height=3, wrap="word")
        q_entry.insert("1.0", question.get("text", ""))
        q_entry.pack(fill="x", padx=30, ipady=4)
        
        tk.Label(win, text="选项（每行一个）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        opt_entry = tk.Text(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                           insertbackground="#eee", bd=0, highlightthickness=1,
                           highlightbackground="#0f3460", height=5, wrap="word")
        if qtype == "judge":
            correct_ans = [chr(65+i) for i, o in enumerate(opts) if o.get("correct")]
            correct_ans = [chr(65+i) for i, o in enumerate(opts) if o.get('correct')]
            opt_txt = chr(9).join(['正确' + ('[正确]' if 'A' in correct_ans else ''), '错误' + ('[正确]' if 'B' in correct_ans else '')])
        else:
            opt_txt = '\n'.join([o.get('text', '') for o in opts])
        opt_entry.insert("1.0", opt_txt)
        opt_entry.pack(fill="x", padx=30, ipady=4)
        
        tk.Label(win, text="答案（多选用ABC，判断A等于正确 B等于错误）", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        correct = [chr(65+i) for i, o in enumerate(opts) if o.get("correct")]
        ans_entry = tk.Entry(win, font=("Microsoft YaHei", 12), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        ans_entry.insert(0, ",".join(correct))
        ans_entry.pack(fill="x", padx=30, ipady=5)
        
        tk.Label(win, text="解析", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#888").pack(pady=(6, 2))
        exp_entry = tk.Entry(win, font=("Microsoft YaHei", 11), bg="#16213e", fg="#eee",
                            insertbackground="#eee", bd=0, highlightthickness=1,
                            highlightbackground="#0f3460")
        exp_entry.insert(0, question.get("explanation", ""))
        exp_entry.pack(fill="x", padx=30, ipady=4)
        
        msg = tk.Label(win, text="", font=("Microsoft YaHei", 9), bg="#1a1a2e", fg="#e94560")
        msg.pack(pady=(4, 2))
        
        def do_save():
            qtext = q_entry.get("1.0", "end-1c").strip()
            new_type = qtype_var.get()
            ans = ans_entry.get().strip()
            exp = exp_entry.get().strip()
            if not qtext:
                msg.config(text="请输入题目")
                return
            if not ans:
                msg.config(text="请输入答案")
                return
            
            bank = load_bank(bank_name)
            if new_type == "judge":
                new_opts = [{"text": "正确", "correct": ans.upper() == "A"},
                           {"text": "错误", "correct": ans.upper() == "B"}]
                bank[idx] = {"text": qtext, "category": "判断题", "type": "judge",
                            "options": new_opts, "explanation": exp,
                            "done": False, "correct": False,
                            "id": bank[idx].get("id", int(random.random() * 10**12))}
            else:
                opt_lines = [l.strip() for l in opt_entry.get("1.0", "end-1c").strip().split(" ") if l.strip()]
                if len(opt_lines) < 2:
                    msg.config(text="至少输入2个选项")
                    return
                new_opts = [{"text": o, "correct": False} for o in opt_lines]
                ans_upper = ans.upper().replace(",", "")
                for c in ans_upper:
                    if c.isalpha() and ord(c) - 65 < len(new_opts):
                        new_opts[ord(c) - 65]["correct"] = True
                if not any(o["correct"] for o in new_opts):
                    msg.config(text="答案不匹配")
                    return
                cat = "多选题" if new_type == "multiple" else "单选题"
                bank[idx] = {"text": qtext, "category": cat, "type": new_type,
                            "options": new_opts, "explanation": exp,
                            "done": False, "correct": False,
                            "id": bank[idx].get("id", int(random.random() * 10**12))}
            save_bank(bank_name, bank)
            messagebox.showinfo("成功", "题目已更新")
            win.destroy()
            self.show_bank_questions(bank_name)
        
        tk.Button(win, text=" 保 存 ", font=("Microsoft YaHei", 12, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=20, pady=6,
                 cursor="hand2", relief="flat", command=do_save).pack(pady=(10, 6))
        
        # 删除按钮
        def do_delete():
            if messagebox.askyesno("确认删除", f"确定删除第{idx+1}题吗？"):
                bank = load_bank(bank_name)
                bank.pop(idx)
                save_bank(bank_name, bank)
                win.destroy()
                self.show_bank_questions(bank_name)
                self.update_top_stats()
        
        tk.Button(win, text="删除此题", font=("Microsoft YaHei", 10),
                 bg="#e94560", fg="#fff", bd=0, padx=12, pady=4,
                 cursor="hand2", relief="flat", command=do_delete).pack()

        def do_save():
            qtext = q_entry.get("1.0", "end-1c").strip()
            new_type = qtype_var.get()
            ans = ans_entry.get().strip()
            exp = exp_entry.get().strip()
            if not qtext:
                msg.config(text="请输入题目")
                return
            if not ans:
                msg.config(text="请输入答案")
                return
            
            bank = load_bank(bank_name)
            if new_type == "judge":
                new_opts = [{"text": "正确", "correct": ans.upper() == "A"},
                           {"text": "错误", "correct": ans.upper() == "B"}]
                bank[idx] = {"text": qtext, "category": "判断题", "type": "judge",
                            "options": new_opts, "explanation": exp,
                            "done": False, "correct": False,
                            "id": bank[idx].get("id", int(random.random() * 10**12))}
            else:
                opt_lines = [l.strip() for l in opt_entry.get("1.0", "end-1c").strip().split("\n") if l.strip()]
                if len(opt_lines) < 2:
                    msg.config(text="至少输入2个选项")
                    return
                new_opts = [{"text": o, "correct": False} for o in opt_lines]
                ans_upper = ans.upper().replace(",", "")
                for c in ans_upper:
                    if c.isalpha() and ord(c) - 65 < len(new_opts):
                        new_opts[ord(c) - 65]["correct"] = True
                if not any(o["correct"] for o in new_opts):
                    msg.config(text="答案不匹配")
                    return
                cat = "多选题" if new_type == "multiple" else "单选题"
                bank[idx] = {"text": qtext, "category": cat, "type": new_type,
                            "options": new_opts, "explanation": exp,
                            "done": False, "correct": False,
                            "id": bank[idx].get("id", int(random.random() * 10**12))}
            save_bank(bank_name, bank)
            messagebox.showinfo("成功", "题目已更新")
            win.destroy()
            self.show_bank_questions(bank_name)
        
        tk.Button(win, text=" 保 存 ", font=("Microsoft YaHei", 12, "bold"),
                 bg="#4ecca3", fg="#1a1a2e", bd=0, padx=20, pady=6,
                 cursor="hand2", relief="flat", command=do_save).pack(pady=(10, 6))
        
        # 删除按钮
    def start_bank_at(self, bank_name, start_idx=0):
        """从指定序号开始练习"""
        self.current_bank = bank_name
        self.current_type_filter = None
        qs = self.get_all_questions()
        filtered = [q for q in qs if not self.current_type_filter or q.get("type") == self.current_type_filter]
        if start_idx >= len(filtered):
            start_idx = 0
        queue = list(filtered)
        self.practice_queue = queue[start_idx:] + queue[:start_idx]
        self.practice_idx = 0
        self.show_practice_question()
if __name__ == "__main__":
    import traceback
    try:
        LoginWindow()
    except Exception as e:
        with open("D:\\刷题软件\\error_log.txt", "w", encoding="utf-8") as f:
            f.write(traceback.format_exc())
        import tkinter.messagebox as mb
        try:
            mb.showerror("错误", f"程序启动失败:\n{e}\n\n错误详情已保存到 error_log.txt")
        except:
            pass